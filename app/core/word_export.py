"""
word_export.py — Genera un .docx judicial fiel al formato real de la Sala Penal.

Formato replicado de EXP N° 244-2026-43 (medido párrafo a párrafo):
  Fuente:        Arial Narrow 11pt en todo el documento
  Página:        A4 (21.001 × 29.700 cm)
  Márgenes:      top 2.251 | left 3.501 | right 1.998 | bottom 2.499
  Interlineado:  1.15 múltiplo
  sp_before/after: 6pt en párrafos principales; 0 en subapartados numerados

  Cabecera institucional:   CENTER, bold
  Bloque metadatos:         JUSTIFY, bold, tab a pos 0.847cm
  Separador línea blanca:   párrafo vacío
  AUTO DE VISTA:            CENTER, bold, underline
  RESOLUCIÓN Nº:            JUSTIFY, bold, underline
  Fecha:                    JUSTIFY, normal
  AUTOS y VISTOS:           JUSTIFY, first_line_indent=1.2488cm (sangría primera línea)
  Secciones I.- II.-:       JUSTIFY, bold, ind_l=0.7514 ind_f=-0.7514
  Párrafo intro sección:    JUSTIFY, ind_l=0.7514, normal (sin sangría first_line)
  Subapartados 6.1.1.:      JUSTIFY, bold+normal, ind_l=0.7514 ind_f=-0.7514, sp_before/after=6
  Citas en cursiva:         JUSTIFY, italic, ind_l=1.0001
  Listas evidencia:         JUSTIFY, ind_l=3.0004 ind_f=-0.4992, bold+underline el título
  DECISIÓN VIII.-:          bold, ind_l=0.7514 ind_f=-0.7514
  Párrafos INFUNDADO/CONFIRMARON: ind_l=1.5011 ind_f=-0.5009, bold keyword + normal
  Cabecera gráfica 1ª pág.: logo + Corte/Sala en encabezado distinto (si hay PNG en resources).
"""

from __future__ import annotations

import os
import re
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from PIL import Image


# ── Constantes exactas medidas del documento real ────────────────────────────

FONT_NAME   = "Arial Narrow"
FONT_SIZE   = Pt(11)

PAGE_WIDTH  = Cm(21.001)
PAGE_HEIGHT = Cm(29.700)
MARGIN_TOP    = Cm(2.251)
MARGIN_LEFT   = Cm(3.501)
MARGIN_RIGHT  = Cm(1.998)
MARGIN_BOTTOM = Cm(2.499)

# Ancho útil entre márgenes (cuerpo / cabecera gráfica)
CONTENT_WIDTH = Cm(21.001 - 3.501 - 1.998)

LINE_SPACING  = 1.15        # múltiplo
SP6 = Pt(6)                 # space_before / space_after estándar
SP0 = Pt(0)

# Indentaciones (medidas exactas del doc real)
IND_SECTION   = Cm(0.7514)   # secciones I.- II.- y subapartados 6.x.x
IND_FIRST_NEG = Cm(-0.7514)  # hanging indent para secciones
IND_BODY_INTRO = Cm(0.7514)  # párrafo intro sin hanging (post-sección)
IND_QUOTE      = Cm(1.0001)  # citas en cursiva
IND_LIST_L     = Cm(3.0004)  # listas de evidencias
IND_LIST_F     = Cm(-0.4992) # hanging de listas
IND_DECISION_L = Cm(1.5011)  # párrafos INFUNDADO/CONFIRMARON
IND_DECISION_F = Cm(-0.5009)
IND_FIRST_BODY = Cm(1.2488)  # primera línea AUTOS y VISTOS
TAB_META       = Cm(0.847)   # tab en bloque de metadatos

_RESOURCES_DIR = Path(__file__).resolve().parent.parent / "resources"
_INSTITUTIONAL_BANNER = _RESOURCES_DIR / "institutional_header_banner.png"
_COL_LOGO = Cm(2.85)   # columna del logotipo en cabecera de 1ª página


# ── Helpers base ─────────────────────────────────────────────────────────────

def _set_line_spacing(para, multiple: float) -> None:
    para.paragraph_format.line_spacing = multiple


def _fmt(para, *,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         sp_before=SP6, sp_after=SP6,
         ind_l=None, ind_f=None, ind_r=None) -> None:
    pf = para.paragraph_format
    pf.alignment    = align
    pf.space_before = sp_before
    pf.space_after  = sp_after
    _set_line_spacing(para, LINE_SPACING)
    if ind_l is not None:
        pf.left_indent = ind_l
    if ind_f is not None:
        pf.first_line_indent = ind_f
    if ind_r is not None:
        pf.right_indent = ind_r


def _run(para, text: str, *, bold=False, italic=False, underline=False) -> Any:
    r = para.add_run(text)
    r.font.name  = FONT_NAME
    r.font.size  = FONT_SIZE
    r.bold       = bold or None      # None = hereda; True = fuerza
    r.italic     = italic or None
    r.underline  = underline or None
    return r


# Regex para **negrita** inline en Markdown
_RE_BOLD   = re.compile(r"\*\*(.*?)\*\*")
_RE_ITALIC = re.compile(r"\*(.*?)\*")   # solo * simple (no doble)


def _inline(para, text: str, *, base_bold=False, base_italic=False) -> None:
    """Renderiza texto con **negrita** y *cursiva* inline."""
    # Primero negrita
    parts = _RE_BOLD.split(text)
    for i, part in enumerate(parts):
        if not part:
            continue
        is_bold = base_bold or (i % 2 == 1)
        # Dentro de cada parte, procesar cursiva
        subparts = _RE_ITALIC.split(part)
        for j, sub in enumerate(subparts):
            if not sub:
                continue
            is_italic = base_italic or (j % 2 == 1)
            _run(para, sub, bold=is_bold, italic=is_italic)


def _institutional_graphic_available() -> bool:
    return _INSTITUTIONAL_BANNER.is_file()


def _crop_logo_tempfile(src: Path) -> Path:
    """Recorta la zona izquierda (logotipo) de la franja ancha; devuelve PNG temporal."""
    im = Image.open(src).convert("RGB")
    w, h = im.size
    if w <= h * 2:
        side = min(w, h)
        logo = im.crop((0, 0, side, h))
    else:
        # Franja tipo banner: cuadrado desde la izquierda usando la altura completa
        side = min(h, max(h, w // 5))
        side = min(side, w)
        logo = im.crop((0, 0, side, h))
    fd, tmp = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    tpath = Path(tmp)
    logo.save(tpath, "PNG")
    return tpath


def _strip_header_body(container) -> None:
    """Elimina párrafos iniciales del encabezado para dejar solo lo que añadamos después."""
    # Story part: remove paragraph XML elements
    for p in list(container.paragraphs):
        el = p._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def _set_cell_shading(cell, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    shd.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shd)


def _set_cell_vertical_align_center(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    val = OxmlElement("w:vAlign")
    val.set(qn("w:val"), "center")
    tc_pr.append(val)


def _run_on_dark(para, text: str, *, bold: bool = True, size_pt: int = 10) -> None:
    for i, line in enumerate(text.split("\n")):
        if i:
            para.add_run().add_break(WD_BREAK.LINE)
        r = para.add_run(line)
        r.font.name = FONT_NAME
        r.font.size = Pt(size_pt)
        r.bold = bold
        r.font.color.rgb = RGBColor(255, 255, 255)


def _build_first_page_institutional_header(doc: Document, meta: dict, tmp_paths: list[Path]) -> None:
    """
    Logo + identificación del órgano solo en la primera página (encabezado distinto).
    Las páginas siguientes usan el encabezado predeterminado vacío.
    """
    corte = meta.get("corte", "CORTE SUPERIOR DE JUSTICIA DE ICA")
    sala = meta.get("sala", "SALA PENAL DE APELACIONES DE CHINCHA Y PISCO")

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Páginas 2+: sin franja
    dh = section.header
    _strip_header_body(dh)

    fh = section.first_page_header
    _strip_header_body(fh)

    tbl = fh.add_table(rows=1, cols=2, width=CONTENT_WIDTH)
    tbl.autofit = False
    tbl.columns[0].width = _COL_LOGO
    tbl.columns[1].width = Cm(CONTENT_WIDTH.cm - _COL_LOGO.cm)

    cell_logo = tbl.cell(0, 0)
    cell_txt = tbl.cell(0, 1)
    _set_cell_vertical_align_center(cell_logo)
    _set_cell_vertical_align_center(cell_txt)
    _set_cell_shading(cell_txt, "000000")

    logo_src = _crop_logo_tempfile(_INSTITUTIONAL_BANNER)
    tmp_paths.append(logo_src)

    p_l = cell_logo.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pic = p_l.add_run()
    run_pic.add_picture(str(logo_src), height=Cm(1.05))

    p_t = cell_txt.paragraphs[0]
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p_t.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    _set_line_spacing(p_t, LINE_SPACING)
    _run_on_dark(p_t, f"{corte}\n{sala}", bold=True, size_pt=10)

    # Encabezado «normal» (resto de páginas): vacío explícito
    if not dh.paragraphs:
        dh.add_paragraph()


# ── Configuración de página y estilos base ───────────────────────────────────

def _page_setup(doc: Document) -> None:
    s = doc.sections[0]
    s.page_width    = PAGE_WIDTH
    s.page_height   = PAGE_HEIGHT
    s.top_margin    = MARGIN_TOP
    s.left_margin   = MARGIN_LEFT
    s.right_margin  = MARGIN_RIGHT
    s.bottom_margin = MARGIN_BOTTOM


def _default_style(doc: Document) -> None:
    """Arial Narrow 11pt como base del documento."""
    style = doc.styles["Normal"]
    style.font.name  = FONT_NAME
    style.font.size  = FONT_SIZE
    style.font.bold  = False


# ── Construcción de la cabecera institucional ────────────────────────────────

def _build_header(doc: Document, meta: dict, *, skip_corte_sala: bool = False) -> None:
    """
    Replica exacta de la cabecera del expediente real:
      párrafo vacío
      CORTE  (CENTER, bold)  — omitido si skip_corte_sala (ya van en cabecera gráfica 1ª pág.)
      SALA   (CENTER, bold)
      párrafo vacío (bold, justify)
      párrafo vacío
      EXPEDIENTE Nº → valor  (JUSTIFY, bold, tab)
      IMPUTADO      → valor  (JUSTIFY, bold, tab, ind_l=4.23 ind_f=-4.23)
      DELITO        → valor  (JUSTIFY, bold, tab)
      AGRAVIADO     → valor  (JUSTIFY, bold, tab)
      PROCEDENCIA   → valor  (JUSTIFY, bold, tab)
      separador
      AUTO DE VISTA  (CENTER, bold, underline)
    """
    corte = meta.get("corte", "CORTE SUPERIOR DE JUSTICIA DE ICA")
    sala  = meta.get("sala",  "SALA PENAL DE APELACIONES DE CHINCHA Y PISCO")

    # Párrafo vacío inicial
    p0 = doc.add_paragraph()
    _fmt(p0)

    if not skip_corte_sala:
        # Corte
        p1 = doc.add_paragraph()
        _fmt(p1, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p1, corte, bold=True)

        # Sala
        p2 = doc.add_paragraph()
        _fmt(p2, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(p2, sala, bold=True)

    # Párrafo vacío bold justify (igual al doc original [003])
    p3 = doc.add_paragraph()
    _fmt(p3)
    _run(p3, "", bold=True)

    # Párrafo vacío [004]
    p4 = doc.add_paragraph()
    _fmt(p4)

    # Bloque metadatos — solo los que tienen valor
    fields = [
        ("EXPEDIENTE Nº", meta.get("expediente", ""), False),
        ("IMPUTADO",      meta.get("imputado",   ""), True),   # True = hanging indent especial
        ("DELITO",        meta.get("delito",     ""), False),
        ("AGRAVIADO",     meta.get("agraviado",  ""), False),
        ("PROCEDENCIA",   meta.get("procedencia","JUZGADO DE INVESTIGACION PREPARATORIA"), False),
    ]
    for label, value, hanging in fields:
        if not value and label not in ("EXPEDIENTE Nº", "PROCEDENCIA"):
            continue
        p = doc.add_paragraph()
        _fmt(p)
        if hanging:
            # Igual a párrafo [006]: ind_l=4.2333 ind_f=-4.2333
            p.paragraph_format.left_indent       = Cm(4.2333)
            p.paragraph_format.first_line_indent = Cm(-4.2333)
        _run(p, f"{label}\t: {value}", bold=True)

    # Separador con espacios (igual a [010])
    p_sep = doc.add_paragraph()
    _fmt(p_sep)
    _run(p_sep, "                                                 \t", bold=True)


# ── Detección de tipo de línea ───────────────────────────────────────────────

# I.- / II.- / III.- etc.  (secciones romanas)
_RE_ROMAN = re.compile(
    r"^([IVXivx]+\.\-\s+|[IVXivx]+\.—\s+)(.*)", re.IGNORECASE
)
# Subapartados: 6.1. / 6.1.1. / 3.1. etc.
_RE_SUBNUM = re.compile(r"^(\d+(?:\.\d+)+\.?\s+)(.*)")
# Decisión: INFUNDADO / CONFIRMARON / DISPUSIERON / REVOCARON / MODIFICARON
_RE_DECISION_KW = re.compile(
    r"^\s*(INFUNDADO|FUNDADO|CONFIRMARON|CONFIRMARON:|REVOCARON|MODIFICARON|DISPUSIERON)\b",
    re.IGNORECASE,
)
# Magistrados (S.S. o solo apellidos en mayúsculas)
_RE_SS = re.compile(r"^S\.S\.?\s*$", re.IGNORECASE)
_RE_MAGISTRADO = re.compile(r"^[A-ZÁÉÍÓÚÑ\s]{5,}$")

# AUTOS y VISTOS  (inicio del cuerpo)
_RE_AUTOS = re.compile(r"^AUTOS\s+Y\s+VISTOS", re.IGNORECASE)

# AUTO DE VISTA / SENTENCIA (título central)
_RE_TITULO = re.compile(
    r"^\s*(AUTO DE VISTA|AUTO DE APELACION|SENTENCIA DE VISTA|SENTENCIA|RESOLUCIÓN|RESOLUCION)\b.*",
    re.IGNORECASE,
)
# RESOLUCIÓN Nº (numeración)
_RE_RESOL_NUM = re.compile(r"^RESOLUCIÓN\s+N[ºo°]\s*\d+", re.IGNORECASE)

# Cita entre comillas (párrafo de cita doctrinal o legal)
_RE_QUOTE = re.compile(r'^[\"\u201c\u201d\u2018\u2019\(…\)].{0,10}')

# Encabezados Markdown
_RE_MD_H12 = re.compile(r"^#{1,2}\s+(.*)")
_RE_MD_H34 = re.compile(r"^#{3,4}\s+(.*)")

# DECISIÓN / VIII.-
_RE_DECISION_SEC = re.compile(
    r"^([IVXivx]+\.\-\s*)?(DECISIÓN|DECISION|FALLO)\b.*", re.IGNORECASE
)

# Listas de evidencias: empieza con guión largo o viñeta o texto bold+underline típico
_RE_EVIDENCIA = re.compile(r"^(Denuncia|Acta|Declaración|Reporte|Parte|Informe|Certificado)\b")


# ── Constructor de párrafos ──────────────────────────────────────────────────

def _add_para_section(doc: Document, text: str) -> None:
    """Secciones I.- II.- III.- (bold, ind_l + hanging)."""
    p = doc.add_paragraph()
    _fmt(p, ind_l=IND_SECTION, ind_f=IND_FIRST_NEG)
    _run(p, text, bold=True)


def _add_para_subnum(doc: Document, num: str, rest: str) -> None:
    """Subapartados 6.1.1. (número bold + texto normal, ind_l + hanging)."""
    p = doc.add_paragraph()
    _fmt(p, ind_l=IND_SECTION, ind_f=IND_FIRST_NEG)
    _run(p, num, bold=True)
    _inline(p, rest)


def _add_para_body(doc: Document, text: str, *, indent_l=None, first_line=None) -> None:
    """Párrafo de cuerpo estándar (justify, 6/6pt)."""
    p = doc.add_paragraph()
    _fmt(p, ind_l=indent_l, ind_f=first_line)
    _inline(p, text)


def _add_para_bold_center(doc: Document, text: str, underline=False) -> None:
    p = doc.add_paragraph()
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, text, bold=True, underline=underline)


def _add_para_bold_justify(doc: Document, text: str, underline=False,
                           ind_l=None, ind_f=None) -> None:
    p = doc.add_paragraph()
    _fmt(p, ind_l=ind_l, ind_f=ind_f)
    _run(p, text, bold=True, underline=underline)


def _add_para_quote(doc: Document, text: str) -> None:
    """Citas en cursiva con indentación."""
    p = doc.add_paragraph()
    _fmt(p, ind_l=IND_QUOTE)
    _run(p, text, italic=True)


def _add_para_autos(doc: Document, text: str) -> None:
    """AUTOS y VISTOS — primera línea indentada (ind_f=1.2488cm)."""
    p = doc.add_paragraph()
    _fmt(p, ind_f=IND_FIRST_BODY)
    # "AUTOS y VISTOS" en bold, resto normal
    m = re.match(r"^(AUTOS\s+[Yy]\s+VISTOS[;,]?\s*)(.*)", text, re.IGNORECASE | re.DOTALL)
    if m:
        _run(p, m.group(1), bold=True)
        _inline(p, m.group(2))
    else:
        _inline(p, text, base_bold=True)


def _add_para_decision_item(doc: Document, text: str) -> None:
    """INFUNDADO / CONFIRMARON / DISPUSIERON — ind_l=1.5011 ind_f=-0.5009."""
    p = doc.add_paragraph()
    _fmt(p, ind_l=IND_DECISION_L, ind_f=IND_DECISION_F)
    # Keyword en bold, resto normal
    m = _RE_DECISION_KW.match(text)
    if m:
        kw = m.group(1)
        rest = text[len(kw):]
        _run(p, "  " + kw, bold=True)
        _inline(p, rest)
    else:
        _inline(p, text.strip())


def _add_para_empty(doc: Document) -> None:
    p = doc.add_paragraph()
    _fmt(p)


# ── Parser principal ─────────────────────────────────────────────────────────

def _parse_lines(doc: Document, lines: list[str]) -> None:
    for raw in lines:
        line = raw.rstrip()

        # Filtrar metadatos internos del .md
        stripped = line.strip()
        if stripped.startswith("<!--") or stripped.startswith("---"):
            continue

        # Línea vacía
        if not stripped:
            _add_para_empty(doc)
            continue

        # Encabezados Markdown → secciones bold
        m = _RE_MD_H12.match(line)
        if m:
            _add_para_section(doc, m.group(1).strip())
            continue
        m = _RE_MD_H34.match(line)
        if m:
            _add_para_bold_justify(doc, m.group(1).strip(),
                                   ind_l=IND_SECTION, ind_f=IND_FIRST_NEG)
            continue

        # AUTO DE VISTA / SENTENCIA (título central)
        if _RE_TITULO.match(stripped) and not _RE_RESOL_NUM.match(stripped):
            _add_para_bold_center(doc, stripped, underline=True)
            continue

        # RESOLUCIÓN Nº 05 (justify, bold, underline)
        if _RE_RESOL_NUM.match(stripped):
            _add_para_bold_justify(doc, stripped, underline=True)
            continue

        # AUTOS y VISTOS
        if _RE_AUTOS.match(stripped):
            _add_para_autos(doc, stripped)
            continue

        # Secciones romanas I.- II.-
        m = _RE_ROMAN.match(stripped)
        if m:
            _add_para_section(doc, m.group(1) + m.group(2))
            continue

        # DECISIÓN / VIII.-
        if _RE_DECISION_SEC.match(stripped):
            _add_para_section(doc, stripped)
            continue

        # Subapartados numerados 6.1. / 6.1.1.
        m = _RE_SUBNUM.match(stripped)
        if m:
            _add_para_subnum(doc, m.group(1), m.group(2))
            continue

        # Items de decisión: INFUNDADO / CONFIRMARON / DISPUSIERON
        if _RE_DECISION_KW.match(stripped):
            _add_para_decision_item(doc, stripped)
            continue

        # S.S. y magistrados
        if _RE_SS.match(stripped):
            _add_para_bold_justify(doc, stripped,
                                   ind_l=IND_SECTION, ind_f=IND_FIRST_NEG)
            continue
        if _RE_MAGISTRADO.match(stripped) and len(stripped.split()) <= 4:
            p = doc.add_paragraph()
            _fmt(p, ind_l=IND_SECTION, ind_f=IND_FIRST_NEG)
            _inline(p, stripped)
            continue

        # Citas en cursiva (empieza con comilla o paréntesis)
        if stripped.startswith(('"', '"', '"', '(…)', "'")) or \
           (stripped.startswith('"') and len(stripped) > 10):
            _add_para_quote(doc, stripped)
            continue

        # Tab inicial → sangría primera línea (párrafo intro post-sección)
        if line.startswith("\t") or line.startswith("    "):
            _add_para_body(doc, stripped, indent_l=IND_BODY_INTRO)
            continue

        # Párrafo de cuerpo normal
        _add_para_body(doc, stripped)


# ── API pública ──────────────────────────────────────────────────────────────

def markdown_to_docx(
    texto_md: str,
    dest_path: Path,
    *,
    metadata: dict | None = None,
    include_header: bool = True,
) -> Path:
    """
    Convierte el texto de Claude a un .docx fiel al formato judicial real.

    Args:
        texto_md:       Texto Markdown/plano de Claude.
        dest_path:      Ruta destino del .docx.
        metadata:       Dict con: expediente, imputado, delito, agraviado,
                        procedencia, corte, sala.
        include_header: Si True, inserta cabecera institucional y, si existe
                        `app/resources/institutional_header_banner.png`, el bloque
                        gráfico (logo + órgano) solo en la 1ª página.

    Returns:
        dest_path confirmado.
    """
    meta = metadata or {}
    doc = Document()
    tmp_cleanup: list[Path] = []

    _default_style(doc)
    _page_setup(doc)

    use_banner = bool(include_header and _institutional_graphic_available())
    if use_banner:
        _build_first_page_institutional_header(doc, meta, tmp_cleanup)

    try:
        if include_header and any(meta.get(k) for k in ("expediente", "imputado", "corte")):
            _build_header(doc, meta, skip_corte_sala=use_banner)

        lines = texto_md.splitlines()
        _parse_lines(doc, lines)

        dest_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(dest_path))
    finally:
        for p in tmp_cleanup:
            try:
                p.unlink(missing_ok=True)
            except OSError:
                pass

    return dest_path


def resolution_docx_filename(materia: str, folder_rel: str, expediente: str = "") -> str:
    safe_exp = re.sub(r"[^\w\-]", "_", expediente).strip("_") if expediente else ""
    safe_mat = re.sub(r"[^\w]", "_", materia)
    if safe_exp:
        return f"EXP_{safe_exp}_{safe_mat}.docx"
    safe_folder = re.sub(r"[^\w\-]", "_", Path(folder_rel).name).strip("_")
    return f"{safe_folder}.docx"
