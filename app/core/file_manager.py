"""
File-system operations for Adiutor Iudicis.
No AI, no API — pure file management.
"""

from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path

from app.core.human_editor_hints import BLOQUE_ROL_ANTIAI_UNA_LINEA_ES, CARGOS_MP_LITERAL_ES
from app.core.prompt_injection_guard import rules_for_untrusted_sources

BASE_DIR = Path(__file__).parent.parent.parent

DIRS = {
    "plantillas":   BASE_DIR / "01_raw" / "plantillas",
    "bibliografia": BASE_DIR / "01_raw" / "bibliografia",
    "resoluciones": BASE_DIR / "03_outputs" / "resoluciones",
    "exports":      BASE_DIR / "03_outputs" / "exports",
    "casos_wiki":   BASE_DIR / "02_wiki" / "casos",
    "instrucciones_generales": BASE_DIR / "01_raw" / "instrucciones_generales",
}

for _d in DIRS.values():
    _d.mkdir(parents=True, exist_ok=True)

# Bibliografía: PDF/Word + texto y Markdown (mejor para Cursor y sin depender de OCR)
BIBLIOGRAFIA_SUFFIXES = frozenset({".pdf", ".doc", ".docx", ".md", ".txt"})

# Archivos que la biblioteca de plantillas muestra por materia (cualquier subcarpeta).
# Incluye formatos que add_plantilla deja sin convertir (.pages, .doc antiguo, etc.).
PLANTILLA_FILE_SUFFIXES = frozenset({
    ".md",
    ".pdf",
    ".doc",
    ".docx",
    ".pages",
    ".txt",
    ".rtf",
    ".odt",
    ".docm",
})

# Subcarpetas por materia del caso (misma clave que la barra lateral de Adiutor Iudicis)
MATERIA_PRISION_PREVENTIVA = "prision_preventiva"
MATERIA_CESACION_PP = "cesacion_prision_preventiva"
MATERIA_PROLONGACION_PP = "prolongacion_prision_preventiva"
MATERIA_BENEFICIOS_PENIT = "beneficios_penitenciarios"
MATERIA_APELACION_SENT = "apelacion_sentencias"
MATERIA_SOBRESEIMIENTO = "apelacion_sobreseimiento"
MATERIA_ENJUICIAMIENTO = "apelacion_enjuiciamiento"
MATERIA_MEDIDAS_COERC = "apelacion_medidas_coercitivas"
MATERIA_TUTELA = "tutela_de_derechos"
MATERIA_NULIDAD = "nulidad_de_actuados"
MATERIA_QUEJAS_DERECHO = "quejas_de_derecho"
MATERIA_RECURSOS_QUEJA = "recursos_de_queja"
MATERIA_CONSULTAS = "consultas"
MATERIA_OTROS = "otros"

MATERIA_SLUGS: frozenset[str] = frozenset(
    {
        MATERIA_PRISION_PREVENTIVA,
        MATERIA_CESACION_PP,
        MATERIA_PROLONGACION_PP,
        MATERIA_BENEFICIOS_PENIT,
        MATERIA_APELACION_SENT,
        MATERIA_SOBRESEIMIENTO,
        MATERIA_ENJUICIAMIENTO,
        MATERIA_MEDIDAS_COERC,
        MATERIA_TUTELA,
        MATERIA_NULIDAD,
        MATERIA_QUEJAS_DERECHO,
        MATERIA_RECURSOS_QUEJA,
        MATERIA_CONSULTAS,
        MATERIA_OTROS,
    }
)
DEFAULT_MATERIA = MATERIA_PRISION_PREVENTIVA

# Etiqueta legible de cada materia (para la UI)
MATERIA_LABELS: dict[str, str] = {
    MATERIA_PRISION_PREVENTIVA: "Prisión preventiva",
    MATERIA_CESACION_PP: "Cesación de prisión preventiva",
    MATERIA_PROLONGACION_PP: "Prolongación de prisión preventiva",
    MATERIA_BENEFICIOS_PENIT: "Beneficios penitenciarios",
    MATERIA_APELACION_SENT: "Apelación de sentencias",
    MATERIA_SOBRESEIMIENTO: "Apelación de autos de sobreseimiento",
    MATERIA_ENJUICIAMIENTO: "Apelación de autos de enjuiciamiento",
    MATERIA_MEDIDAS_COERC: "Apelación de medidas coercitivas",
    MATERIA_TUTELA: "Tutela de derechos",
    MATERIA_NULIDAD: "Nulidad de actuados",
    MATERIA_QUEJAS_DERECHO: "Quejas de derecho",
    MATERIA_RECURSOS_QUEJA: "Recursos de queja",
    MATERIA_CONSULTAS: "Consultas",
    MATERIA_OTROS: "Otros",
}

# Etiquetas de los 6 slots de documentos por materia. Estructura universal:
#   solicitud_inicial → lo que una parte le pidió al juez de 1ra instancia
#   resolucion_apelada → lo que el juez de 1ra instancia decidió sobre esa solicitud
#   recurso_apelacion → el escrito de apelación que llega a la Sala
#   anexos → actos de investigación / pruebas / documentos adjuntos
#   audio → audio o transcripción de audiencia de apelación
#   otros → cajón libre para lo que no encaje en los demás
SLOT_KEYS: tuple[str, ...] = (
    "solicitud_inicial",
    "resolucion_apelada",
    "recurso_apelacion",
    "anexos",
    "audio",
    "otros",
)

_DEFAULT_SLOT_LABELS: dict[str, str] = {
    "solicitud_inicial": "Solicitud / requerimiento inicial",
    "resolucion_apelada": "Resolución de 1ra instancia (apelada)",
    "recurso_apelacion": "Recurso de apelación",
    "anexos": "Anexos / actos de investigación",
    "audio": "Audio de audiencia",
    "otros": "Otros",
}

SLOT_LABELS: dict[str, dict[str, str]] = {
    MATERIA_PRISION_PREVENTIVA: {
        "solicitud_inicial": "Requerimiento fiscal de prisión preventiva",
        "resolucion_apelada": "Auto que concede/deniega la prisión preventiva",
        "recurso_apelacion": "Recurso de apelación",
        "anexos": "Anexos / actos de investigación",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_CESACION_PP: {
        "solicitud_inicial": "Solicitud de cesación de PP",
        "resolucion_apelada": "Auto que resuelve la cesación",
        "recurso_apelacion": "Recurso de apelación",
        "anexos": "Anexos / nuevos elementos",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_PROLONGACION_PP: {
        "solicitud_inicial": "Requerimiento fiscal de prolongación",
        "resolucion_apelada": "Auto que concede/deniega la prolongación",
        "recurso_apelacion": "Recurso de apelación",
        "anexos": "Anexos / actos de investigación",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_BENEFICIOS_PENIT: {
        "solicitud_inicial": "Solicitud del interno",
        "resolucion_apelada": "Resolución que concede/deniega el beneficio",
        "recurso_apelacion": "Recurso de apelación",
        "anexos": "Informes penitenciarios y anexos",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_APELACION_SENT: {
        "solicitud_inicial": "Acusación fiscal",
        "resolucion_apelada": "Sentencia de 1ra instancia",
        "recurso_apelacion": "Recurso de apelación",
        "anexos": "Actuados del juicio / elementos probatorios",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_SOBRESEIMIENTO: {
        "solicitud_inicial": "Requerimiento fiscal de sobreseimiento",
        "resolucion_apelada": "Auto de sobreseimiento",
        "recurso_apelacion": "Recurso de apelación",
        "anexos": "Actos de investigación",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_ENJUICIAMIENTO: {
        "solicitud_inicial": "Acusación fiscal",
        "resolucion_apelada": "Auto de enjuiciamiento",
        "recurso_apelacion": "Recurso de apelación",
        "anexos": "Actos de investigación",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_MEDIDAS_COERC: {
        "solicitud_inicial": "Requerimiento fiscal de la medida",
        "resolucion_apelada": "Auto que impone/rechaza la medida",
        "recurso_apelacion": "Recurso de apelación",
        "anexos": "Anexos / elementos",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_TUTELA: {
        "solicitud_inicial": "Solicitud de tutela de derechos",
        "resolucion_apelada": "Resolución que resuelve la tutela",
        "recurso_apelacion": "Recurso de apelación",
        "anexos": "Anexos",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_NULIDAD: {
        "solicitud_inicial": "Solicitud / pedido que originó la resolución",
        "resolucion_apelada": "Resolución cuestionada",
        "recurso_apelacion": "Recurso planteando la nulidad",
        "anexos": "Anexos",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_QUEJAS_DERECHO: {
        "solicitud_inicial": "Pedido denegado",
        "resolucion_apelada": "Resolución que deniega el recurso",
        "recurso_apelacion": "Queja de derecho",
        "anexos": "Anexos",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_RECURSOS_QUEJA: {
        "solicitud_inicial": "Pedido originario",
        "resolucion_apelada": "Resolución que deniega recurso",
        "recurso_apelacion": "Recurso de queja",
        "anexos": "Anexos",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_CONSULTAS: {
        "solicitud_inicial": "Pedido o petición",
        "resolucion_apelada": "Resolución consultada",
        "recurso_apelacion": "Escrito elevatorio",
        "anexos": "Anexos",
        "audio": "Audio de audiencia",
        "otros": "Otros",
    },
    MATERIA_OTROS: dict(_DEFAULT_SLOT_LABELS),
}


def materia_label(materia: str) -> str:
    """Etiqueta legible para mostrar en la UI."""
    return MATERIA_LABELS.get(materia, materia.replace("_", " ").capitalize())


def slot_labels_for(materia: str) -> dict[str, str]:
    """Etiquetas de los 6 slots adaptadas a la materia."""
    return SLOT_LABELS.get(materia, dict(_DEFAULT_SLOT_LABELS))


MIGRATION_FLAG = BASE_DIR / "01_raw" / ".wikijuez_materias_layout_v1"
CASOS_MIGRATION_FLAG = BASE_DIR / "01_raw" / ".wikijuez_casos_por_materia_v1"


def _validate_materia(materia: str) -> str:
    if materia not in MATERIA_SLUGS:
        raise ValueError(f"Materia no válida: {materia!r}")
    return materia


def dir_instrucciones_generales() -> Path:
    """Carpeta `01_raw/instrucciones_generales/` — un .md por materia."""
    d = DIRS["instrucciones_generales"]
    d.mkdir(parents=True, exist_ok=True)
    return d


def path_instruccion_general(materia: str) -> Path:
    _validate_materia(materia)
    return dir_instrucciones_generales() / f"{materia}.md"


def read_instruccion_general(materia: str) -> str:
    """Texto de la instrucción general persistida para la materia (vacío si no existe)."""
    _validate_materia(materia)
    p = path_instruccion_general(materia)
    if not p.is_file():
        return ""
    return p.read_text(encoding="utf-8", errors="replace")


def save_instruccion_general(materia: str, text: str) -> None:
    """Guarda la instrucción general de la materia en disco."""
    _validate_materia(materia)
    path_instruccion_general(materia).write_text(text, encoding="utf-8")


def dir_bibliografia_materia(materia: str = DEFAULT_MATERIA) -> Path:
    _validate_materia(materia)
    d = DIRS["bibliografia"] / materia
    d.mkdir(parents=True, exist_ok=True)
    return d


def dir_plantillas_materia(materia: str = DEFAULT_MATERIA) -> Path:
    _validate_materia(materia)
    d = DIRS["plantillas"] / materia
    d.mkdir(parents=True, exist_ok=True)
    return d


def dir_resoluciones_materia(materia: str = DEFAULT_MATERIA) -> Path:
    _validate_materia(materia)
    d = DIRS["resoluciones"] / materia
    d.mkdir(parents=True, exist_ok=True)
    return d


# ---------------------------------------------------------------------------
# Corpus del magistrado (resoluciones previas ya firmadas por el juez)
# ---------------------------------------------------------------------------

CORPUS_INDEX_NAME = ".wikijuez_corpus_index.json"


def dir_corpus_materia(materia: str = DEFAULT_MATERIA) -> Path:
    """01_raw/<materia>/corpus_magistrado/ — resoluciones previas del juez.

    El magistrado arroja aquí (vía Finder) sus PDFs de resoluciones ya firmadas.
    La app los detecta en el siguiente arranque y propone ingestarlos para
    generar fichas wiki reutilizables como referencia de estilo.
    """
    _validate_materia(materia)
    d = BASE_DIR / "01_raw" / materia / "corpus_magistrado"
    d.mkdir(parents=True, exist_ok=True)
    return d


def dir_casos_previos_wiki(materia: str = DEFAULT_MATERIA) -> Path:
    """02_wiki/casos_previos/<materia>/ — fichas wiki generadas del corpus."""
    _validate_materia(materia)
    d = BASE_DIR / "02_wiki" / "casos_previos" / materia
    d.mkdir(parents=True, exist_ok=True)
    return d


_CORPUS_INGEST_SUFFIXES = frozenset({".pdf", ".docx", ".doc"})


def list_corpus_pdfs(materia: str = DEFAULT_MATERIA) -> list[Path]:
    """PDF, DOCX y DOC en 01_raw/<materia>/corpus_magistrado/ (recursivo).

    El nombre histórico «pdfs» se mantiene por compatibilidad.
    """
    _validate_materia(materia)
    base = dir_corpus_materia(materia)
    out: list[Path] = []
    for f in base.rglob("*"):
        if (
            f.is_file()
            and not f.name.startswith(".")
            and f.suffix.lower() in _CORPUS_INGEST_SUFFIXES
        ):
            out.append(f)
    return sorted(out, key=lambda p: p.name.lower())


def _corpus_index_path(materia: str) -> Path:
    """JSON que registra qué PDFs del corpus ya fueron procesados a wiki."""
    _validate_materia(materia)
    return dir_corpus_materia(materia) / CORPUS_INDEX_NAME


def read_corpus_index(materia: str = DEFAULT_MATERIA) -> dict[str, str]:
    """Devuelve {nombre_pdf: ruta_ficha_wiki} para los ya procesados."""
    import json
    p = _corpus_index_path(materia)
    if not p.is_file():
        return {}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}


def write_corpus_index(materia: str, index: dict[str, str]) -> None:
    """Persiste el índice de PDFs del corpus ya procesados."""
    import json
    _validate_materia(materia)
    p = _corpus_index_path(materia)
    p.write_text(json.dumps(index, indent=2, ensure_ascii=False), encoding="utf-8")


def pending_corpus_pdfs(materia: str = DEFAULT_MATERIA) -> list[Path]:
    """PDFs del corpus que aún no tienen ficha wiki generada."""
    idx = read_corpus_index(materia)
    return [p for p in list_corpus_pdfs(materia) if p.name not in idx]


def list_corpus_wiki_fichas(materia: str = DEFAULT_MATERIA) -> list[Path]:
    """Fichas .md ya generadas en 02_wiki/casos_previos/<materia>/."""
    _validate_materia(materia)
    base = dir_casos_previos_wiki(materia)
    return sorted(
        [
            f for f in base.iterdir()
            if f.is_file()
            and f.suffix.lower() == ".md"
            and not f.name.startswith(".")
        ],
        key=lambda p: p.name.lower(),
    )


# Manifiesto de clasificación de fuentes por ranura (6 slots)
FUENTES_SLOT_MANIFEST = ".wikijuez_slot_manifest.json"


def legacy_infer_slot_paths(case_folder: Path) -> dict[str, list[Path]]:
    """Sin manifiesto: archivos en fuentes/<slot>/… o sueltos en fuentes/ → otros."""
    out: dict[str, list[Path]] = {k: [] for k in SLOT_KEYS}
    fu = case_folder / "fuentes"
    if not fu.is_dir():
        return out
    for f in fu.rglob("*"):
        if not f.is_file():
            continue
        if f.name.startswith(".") or f.name == FUENTES_SLOT_MANIFEST:
            continue
        rel = f.relative_to(fu)
        if not rel.parts:
            continue
        top = rel.parts[0]
        if top in SLOT_KEYS:
            out[top].append(f)
        else:
            out["otros"].append(f)
    return out


def read_fuentes_slots(case_folder: Path) -> dict[str, list[Path]]:
    """Lee manifiesto si existe; si no, infiere por carpetas o sueltos en fuentes/."""
    fu = case_folder / "fuentes"
    man_path = fu / FUENTES_SLOT_MANIFEST
    out: dict[str, list[Path]] = {k: [] for k in SLOT_KEYS}
    if man_path.is_file():
        try:
            data = json.loads(man_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            data = {}
        for k, rels in data.items():
            if k not in SLOT_KEYS:
                continue
            if not isinstance(rels, list):
                continue
            for rel in rels:
                if not isinstance(rel, str):
                    continue
                p = fu / rel
                if p.is_file():
                    out[k].append(p)
        if any(out.values()):
            return out
    return legacy_infer_slot_paths(case_folder)


def write_fuentes_slots_manifest(
    case_folder: Path, slot_to_relpaths: dict[str, list[str]],
) -> None:
    """Persiste rutas relativas a fuentes/ (p. ej. anexos/informe.pdf)."""
    fu = case_folder / "fuentes"
    fu.mkdir(parents=True, exist_ok=True)
    p = fu / FUENTES_SLOT_MANIFEST
    clean = {k: v for k, v in slot_to_relpaths.items() if v}
    if clean:
        p.write_text(json.dumps(clean, indent=2, ensure_ascii=False), encoding="utf-8")
    elif p.exists():
        try:
            p.unlink()
        except OSError:
            pass


def migrate_legacy_materia_folders() -> None:
    """Una sola vez: archivos sueltos en bibliografía/plantillas/resoluciones → prisión preventiva."""
    if MIGRATION_FLAG.exists():
        return
    pp = DEFAULT_MATERIA
    for m in MATERIA_SLUGS:
        (DIRS["bibliografia"] / m).mkdir(parents=True, exist_ok=True)
        (DIRS["plantillas"] / m).mkdir(parents=True, exist_ok=True)
        (DIRS["resoluciones"] / m).mkdir(parents=True, exist_ok=True)
        (DIRS["exports"] / m).mkdir(parents=True, exist_ok=True)

    bib = DIRS["bibliografia"]
    for f in list(bib.iterdir()):
        if f.is_file() and f.suffix.lower() in BIBLIOGRAFIA_SUFFIXES:
            dest = bib / pp / f.name
            if not dest.exists():
                shutil.move(str(f), str(dest))

    res = DIRS["resoluciones"]
    for f in list(res.iterdir()):
        if f.is_file() and f.suffix.lower() == ".md":
            dest = res / pp / f.name
            if not dest.exists():
                shutil.move(str(f), str(dest))

    pla = DIRS["plantillas"]
    for f in list(pla.iterdir()):
        if f.name in MATERIA_SLUGS:
            continue
        dest = pla / pp / f.name
        if not dest.exists():
            shutil.move(str(f), str(dest))

    exp = DIRS["exports"]
    for f in list(exp.iterdir()):
        if f.is_file():
            dest = exp / pp / f.name
            if not dest.exists():
                shutil.move(str(f), str(dest))

    MIGRATION_FLAG.write_text(
        "Migración v1: archivos existentes asociados a prisión preventiva.\n",
        encoding="utf-8",
    )


def migrate_legacy_caso_folders_to_materia() -> None:
    """Una sola vez: 01_raw/caso_NNN_* → 01_raw/prision_preventiva/caso_NNN_*."""
    if CASOS_MIGRATION_FLAG.exists():
        return
    raw = BASE_DIR / "01_raw"
    pp = DEFAULT_MATERIA
    dest_parent = raw / pp
    dest_parent.mkdir(parents=True, exist_ok=True)
    for entry in list(raw.iterdir()):
        if not entry.is_dir():
            continue
        if entry.name.startswith("caso_"):
            dest = dest_parent / entry.name
            if not dest.exists():
                shutil.move(str(entry), str(dest))
    CASOS_MIGRATION_FLAG.write_text(
        "Expedientes en 01_raw/<materia>/caso_NNN_…\n", encoding="utf-8"
    )


def migrate_stray_root_caso_folders() -> None:
    """En cada arranque: mueve cualquier 01_raw/caso_* suelto → prisión preventiva.

    Complementa la migración única: carpetas nuevas o copiadas a la raíz vuelven
    a la ubicación canónica del repositorio histórico (PP).
    """
    raw = BASE_DIR / "01_raw"
    if not raw.is_dir():
        return
    dest_parent = raw / DEFAULT_MATERIA
    dest_parent.mkdir(parents=True, exist_ok=True)
    for entry in list(raw.iterdir()):
        if not entry.is_dir():
            continue
        if not entry.name.startswith("caso_") or entry.name.startswith("."):
            continue
        dest = dest_parent / entry.name
        if not dest.exists():
            shutil.move(str(entry), str(dest))


def _raw_root_caso_dirs() -> list[Path]:
    """Directorios caso_* directamente bajo 01_raw/ (legado previo a subcarpeta materia)."""
    raw = BASE_DIR / "01_raw"
    if not raw.is_dir():
        return []
    return [
        x
        for x in raw.iterdir()
        if x.is_dir() and x.name.startswith("caso_") and not x.name.startswith(".")
    ]


def _dedupe_paths(paths: list[Path]) -> list[Path]:
    seen: set[Path] = set()
    out: list[Path] = []
    for p in paths:
        try:
            k = p.resolve()
        except OSError:
            k = p
        if k not in seen:
            seen.add(k)
            out.append(p)
    return out


def _sort_case_folders_reciente(paths: list[Path]) -> list[Path]:
    """Más reciente primero (mtime de la carpeta)."""

    def mtime(p: Path) -> float:
        try:
            return p.stat().st_mtime
        except OSError:
            return 0.0

    return sorted(paths, key=mtime, reverse=True)


def caso_num_from_folder_key(folder_key: str) -> str:
    """Extrae el número NNN de 'caso_NNN_slug' o 'materia/caso_NNN_slug'."""
    seg = folder_key.replace("\\", "/").strip("/").split("/")[-1]
    parts = seg.split("_")
    if len(parts) >= 2 and seg.startswith("caso_"):
        return parts[1]
    return seg


def infer_materia_from_resoluciones_md(md_path: Path) -> str:
    """Si la ruta está bajo resoluciones/<materia>/, devuelve esa materia (cualquier extensión)."""
    try:
        rel = md_path.resolve().relative_to(DIRS["resoluciones"].resolve())
    except ValueError:
        return DEFAULT_MATERIA
    parts = rel.parts
    if len(parts) >= 1 and parts[0] in MATERIA_SLUGS:
        return parts[0]
    return DEFAULT_MATERIA


# Modo «continuar borrador»: formatos admitidos en el selector y en la validación
BORRADOR_CONTINUAR_SUFFIXES = frozenset({".md", ".doc", ".docx", ".pdf"})
BORRADOR_CONTINUAR_QFILE_FILTER = (
    "Borrador (*.md *.doc *.docx *.pdf *.pages);;"
    "Markdown (*.md);;Word (*.doc *.docx);;PDF (*.pdf);;Pages (*.pages);;Todos los archivos (*)"
)


def is_valid_borrador_continuar_path(path: Path | None) -> bool:
    """True si existe y es .md, .doc, .docx, .pdf o paquete Apple Pages (.pages archivo o carpeta)."""
    if path is None:
        return False
    if not path.exists():
        return False
    suf = path.suffix.lower()
    if path.is_file():
        return suf in BORRADOR_CONTINUAR_SUFFIXES or suf == ".pages"
    if suf == ".pages" and path.is_dir():
        return True
    return False


# ---------------------------------------------------------------------------
# Case numbering
# ---------------------------------------------------------------------------

def _collect_caso_numbers_from_dir(base: Path) -> list[int]:
    numbers: list[int] = []
    if not base.is_dir():
        return numbers
    for entry in base.iterdir():
        if entry.is_dir() and entry.name.startswith("caso_"):
            parts = entry.name.split("_")
            if len(parts) >= 2:
                try:
                    numbers.append(int(parts[1]))
                except ValueError:
                    pass
    return numbers


def get_next_case_number(materia: str = DEFAULT_MATERIA) -> int:
    """Siguiente número de expediente dentro de 01_raw/<materia>/caso_NNN_*."""
    _validate_materia(materia)
    numbers: list[int] = _collect_caso_numbers_from_dir(BASE_DIR / "01_raw" / materia)
    # Repositorio histórico unificado en PP: números en 01_raw/caso_* (raíz) cuentan igual
    if materia == MATERIA_PRISION_PREVENTIVA:
        for entry in _raw_root_caso_dirs():
            parts = entry.name.split("_")
            if len(parts) >= 2:
                try:
                    numbers.append(int(parts[1]))
                except ValueError:
                    pass
    return max(numbers, default=0) + 1


def case_folder_name(numero: int, nombre: str) -> str:
    slug = nombre.strip().lower().replace(" ", "_")
    return f"caso_{numero:03d}_{slug}"


def create_case_folder(
    numero: int, nombre: str, materia: str = DEFAULT_MATERIA,
) -> Path:
    """Crea 01_raw/<materia>/caso_NNN_nombre/ con subcarpeta fuentes/."""
    _validate_materia(materia)
    folder = BASE_DIR / "01_raw" / materia / case_folder_name(numero, nombre)
    folder.mkdir(parents=True, exist_ok=True)
    (folder / "fuentes").mkdir(exist_ok=True)
    return folder


# ---------------------------------------------------------------------------
# File listings
# ---------------------------------------------------------------------------

def _list_files(directory: Path) -> list[Path]:
    if not directory.exists():
        return []
    return sorted(
        [f for f in directory.iterdir() if f.is_file() and not f.name.startswith(".")],
        key=lambda p: p.name.lower(),
    )


def list_plantillas(materia: str = DEFAULT_MATERIA) -> list[Path]:
    """Lista plantillas y documentos de apoyo bajo 01_raw/plantillas/<materia>/ (toda profundidad)."""
    _validate_materia(materia)
    base = DIRS["plantillas"] / materia
    if not base.exists():
        return []
    out: list[Path] = []
    for f in base.rglob("*"):
        if (
            f.is_file()
            and not f.name.startswith(".")
            and f.suffix.lower() in PLANTILLA_FILE_SUFFIXES
        ):
            out.append(f)
    return sorted(out, key=lambda p: str(p).lower())


def list_bibliografia(materia: str = DEFAULT_MATERIA) -> list[Path]:
    """Documentos de apoyo bajo 01_raw/bibliografia/<materia>/ (incluye subcarpetas)."""
    _validate_materia(materia)
    base = DIRS["bibliografia"] / materia
    if not base.is_dir():
        return []
    out: list[Path] = []
    for f in base.rglob("*"):
        if (
            f.is_file()
            and not f.name.startswith(".")
            and f.suffix.lower() in BIBLIOGRAFIA_SUFFIXES
        ):
            out.append(f)
    return sorted(out, key=lambda p: str(p).lower())


def list_case_folders(materia: str | None = None) -> list[Path]:
    """Expedientes en 01_raw/<materia>/caso_*.

    Prisión preventiva incluye además cualquier 01_raw/caso_* en la raíz (legado).
    Orden: más reciente primero (mtime).

    Si materia es None, lista todos los expedientes de las tres materias
    (p. ej. selector global en «Agregar fuentes»).
    """
    raw = BASE_DIR / "01_raw"
    if materia is not None:
        _validate_materia(materia)
        d = raw / materia
        folders: list[Path] = []
        if d.is_dir():
            folders.extend(
                x
                for x in d.iterdir()
                if x.is_dir() and x.name.startswith("caso_") and not x.name.startswith(".")
            )
        if materia == MATERIA_PRISION_PREVENTIVA:
            folders.extend(_raw_root_caso_dirs())
            folders = _dedupe_paths(folders)
        return _sort_case_folders_reciente(folders)
    out: list[Path] = []
    for m in sorted(MATERIA_SLUGS):
        out.extend(list_case_folders(materia=m))
    return _dedupe_paths(_sort_case_folders_reciente(out))


def list_case_files(case_folder: Path) -> list[Path]:
    """Recursively list all files inside a case folder (no hidden files)."""
    return sorted(
        [f for f in case_folder.rglob("*")
         if f.is_file() and not f.name.startswith(".")],
        key=lambda p: (p.parent.name, p.name.lower()),
    )


def add_to_case(src: Path, case_folder: Path, subfolder: str | None = None) -> Path:
    """Copy a file into a case folder, optionally into a subfolder."""
    dest_dir = case_folder / subfolder if subfolder else case_folder
    return copy_to(src, dest_dir / src.name)


def list_resoluciones(materia: str = DEFAULT_MATERIA) -> list[Path]:
    """Solo .md en 03_outputs/resoluciones/<materia>/ — más recientes primero."""
    _validate_materia(materia)
    base = DIRS["resoluciones"] / materia
    if not base.exists():
        return []
    files = [
        f
        for f in base.iterdir()
        if f.is_file() and f.suffix.lower() == ".md" and not f.name.startswith(".")
    ]
    files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return files


def _safe_caso_filename_stem(folder_rel: str) -> str:
    """Último segmento de 01_raw/materia/caso_NNN_slug → nombre de archivo seguro."""
    seg = folder_rel.replace("\\", "/").strip("/").split("/")[-1]
    out: list[str] = []
    for c in seg:
        if c.isalnum() or c in "._-":
            out.append(c)
        else:
            out.append("_")
    s = "".join(out).strip("_")
    return s or "caso"


def find_resolucion_md_for_case(folder: Path, materia: str) -> Path | None:
    """
    Localiza el .md de resolución guardado para una carpeta de expediente (`01_raw/...`).

    El nombre del archivo usa `_safe_caso_filename_stem` (p. ej. espacios → «_»). Si la
    búsqueda usara solo `folder.name` sin sanitizar, el glob no halla el archivo y el
    panel de redacción queda vacío aunque el .md exista en «Resoluciones».
    """
    _validate_materia(materia)
    res_dir = DIRS["resoluciones"] / materia
    if not res_dir.is_dir():
        return None
    raw_root = BASE_DIR / "01_raw"
    try:
        folder_rel = str(folder.relative_to(raw_root)).replace("\\", "/")
    except ValueError:
        folder_rel = folder.name
    stem_safe = _safe_caso_filename_stem(folder_rel)
    for stem in (stem_safe, folder.name):
        if not stem:
            continue
        cands = sorted(
            (f for f in res_dir.glob(f"{stem}_resolucion*.md") if f.is_file()),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        )
        if cands:
            return cands[0]
    marker = f"01_raw/{folder_rel}/"
    for f in sorted(
        (p for p in res_dir.glob("*_resolucion*.md") if p.is_file()),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    ):
        try:
            head = f.read_text(encoding="utf-8", errors="replace")[:4000]
        except OSError:
            continue
        if marker in head:
            return f
    return None


def save_prompt_to_resoluciones_folder(
    materia: str, folder_rel: str, prompt_text: str,
) -> Path:
    """Escribe el prompt en 03_outputs/resoluciones/<materia>/…_PROMPT_para_Cursor.md."""
    _validate_materia(materia)
    base = dir_resoluciones_materia(materia)
    stem = _safe_caso_filename_stem(folder_rel)
    name = f"{stem}_PROMPT_para_Cursor.md"
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    content = (
        f"<!-- Adiutor Iudicis · prompt para Cursor · {ts} · expediente: 01_raw/{folder_rel}/ -->\n\n"
        f"{prompt_text}"
    )
    path = base / name
    path.write_text(content, encoding="utf-8")
    return path


def save_resolucion_cursor_text(
    materia: str, folder_rel: str, texto: str,
) -> Path:
    """Guarda el texto que Cursor devolvió; aparece en la pestaña Resoluciones."""
    _validate_materia(materia)
    base = dir_resoluciones_materia(materia)
    stem = _safe_caso_filename_stem(folder_rel)
    name = f"{stem}_resolucion.md"
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    header = (
        f"<!-- Adiutor Iudicis · resolución (Cursor) · guardado {ts} · 01_raw/{folder_rel}/ -->\n\n"
    )
    path = base / name
    path.write_text(header + texto.strip(), encoding="utf-8")
    return path


def save_resolucion_generada_backup(
    materia: str, folder_rel: str, texto: str, *, sufijo_slug: str,
) -> Path:
    """Copia paralela cuando la generación vino incompleta o diferente por reintentos."""
    _validate_materia(materia)
    safe = "".join(c if (c.isalnum() or c in "._-") else "_" for c in (sufijo_slug or "backup")).strip("_")
    if not safe:
        safe = "backup"
    safe = safe[:140]
    base = dir_resoluciones_materia(materia)
    stem = _safe_caso_filename_stem(folder_rel)
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    name = f"{stem}_resolucion_{safe}_{ts}.md"
    head_ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    header = (
        f"<!-- Adiutor Iudicis · copia backup generación Claude · {head_ts} · "
        f"ref: {stem} · sufijo:{safe} · 01_raw/{folder_rel}/ -->\n\n"
    )
    path = base / name
    path.write_text(header + texto.strip(), encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# Copy helpers
# ---------------------------------------------------------------------------

def copy_to(src: Path, dest: Path) -> Path:
    dest.parent.mkdir(parents=True, exist_ok=True)
    if dest.resolve() != src.resolve():
        shutil.copy2(src, dest)
    return dest


def _docx_to_text(src: Path) -> str:
    """Extract plain text from a .docx file using only stdlib (no python-docx needed)."""
    import zipfile
    import re
    with zipfile.ZipFile(src) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    paras = re.findall(r"<w:p[ >].*?</w:p>", xml, re.DOTALL)
    lines = []
    for p in paras:
        texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", p, re.DOTALL)
        line = "".join(texts).strip()
        if line and not line.startswith("<w:"):
            lines.append(line)
    return "\n\n".join(lines)


def _pdf_to_text(src: Path) -> str:
    """Extract plain text from a PDF (pdfplumber + OCR fallback)."""
    from app.core.pdf_extract import extract_pdf_text

    return extract_pdf_text(src)


def add_plantilla(
    src: Path, tipo_carpeta: str = "", materia: str = DEFAULT_MATERIA,
) -> Path:
    """Copy a plantilla bajo 01_raw/plantillas/<materia>/…, convirtiendo a .md si aplica."""
    _validate_materia(materia)
    dest_dir = DIRS["plantillas"] / materia
    if tipo_carpeta:
        slug = tipo_carpeta.strip().lower().replace(" ", "_")
        dest_dir = dest_dir / slug
    dest_dir.mkdir(parents=True, exist_ok=True)

    ext = src.suffix.lower()
    md_stem = src.stem

    if ext == ".docx":
        text = _docx_to_text(src)
        dest = dest_dir / f"{md_stem}.md"
        dest.write_text(text, encoding="utf-8")
        return dest
    elif ext == ".pdf":
        text = _pdf_to_text(src)
        dest = dest_dir / f"{md_stem}.md"
        dest.write_text(text, encoding="utf-8")
        return dest
    elif ext == ".txt":
        dest = dest_dir / f"{md_stem}.md"
        dest.write_text(src.read_text(encoding="utf-8", errors="replace"), encoding="utf-8")
        return dest
    else:
        # .md or unknown — copy as-is
        return copy_to(src, dest_dir / src.name)


def add_bibliografia(src: Path, materia: str = DEFAULT_MATERIA) -> Path:
    ext = src.suffix.lower()
    if ext not in BIBLIOGRAFIA_SUFFIXES:
        raise ValueError(
            "La bibliografía admite PDF, Word, Markdown y texto plano "
            "(.pdf, .doc, .docx, .md, .txt). "
            f"Archivo: {src.name}"
        )
    dest_dir = dir_bibliografia_materia(materia)
    return copy_to(src, dest_dir / src.name)


def dir_bibliografia_global() -> Path:
    """Bibliografía global (CPP, CP, Constitución) — compartida entre todas las materias."""
    d = DIRS["bibliografia"] / "global"
    d.mkdir(parents=True, exist_ok=True)
    return d


def list_bibliografia_global() -> list[Path]:
    """Archivos en 01_raw/bibliografia/global/."""
    base = dir_bibliografia_global()
    out: list[Path] = []
    for f in base.rglob("*"):
        if (
            f.is_file()
            and not f.name.startswith(".")
            and f.suffix.lower() in BIBLIOGRAFIA_SUFFIXES
        ):
            out.append(f)
    return sorted(out, key=lambda p: p.name.lower())


def add_bibliografia_global(src: Path) -> Path:
    """Copia un archivo a la bibliografía global."""
    ext = src.suffix.lower()
    if ext not in BIBLIOGRAFIA_SUFFIXES:
        raise ValueError(
            "La bibliografía global admite PDF, Word, Markdown y texto plano "
            "(.pdf, .doc, .docx, .md, .txt). "
            f"Archivo: {src.name}"
        )
    return copy_to(src, dir_bibliografia_global() / src.name)


# ---------------------------------------------------------------------------
# Fichas wiki desde bibliografía (doctrina) → 02_wiki/bibliografia/
# ---------------------------------------------------------------------------

BIBLIO_WIKI_INDEX_NAME = ".wikijuez_bibliografia_fichas_index.json"
BIBLIO_GLOBAL_WIKI_INDEX_NAME = ".wikijuez_bibliografia_global_fichas_index.json"


def dir_bibliografia_wiki(materia: str = DEFAULT_MATERIA) -> Path:
    """02_wiki/bibliografia/<materia>/ — fichas generadas desde 01_raw/bibliografia/<materia>/."""
    _validate_materia(materia)
    d = BASE_DIR / "02_wiki" / "bibliografia" / materia
    d.mkdir(parents=True, exist_ok=True)
    return d


def dir_bibliografia_global_wiki() -> Path:
    """02_wiki/bibliografia/global/ — fichas desde 01_raw/bibliografia/global/."""
    d = BASE_DIR / "02_wiki" / "bibliografia" / "global"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _bibliografia_wiki_index_path(materia: str) -> Path:
    _validate_materia(materia)
    return dir_bibliografia_materia(materia) / BIBLIO_WIKI_INDEX_NAME


def _bibliografia_global_wiki_index_path() -> Path:
    return dir_bibliografia_global() / BIBLIO_GLOBAL_WIKI_INDEX_NAME


def read_bibliografia_wiki_index(materia: str = DEFAULT_MATERIA) -> dict[str, str]:
    """{clave_relativa_origen → ruta_ficha_md relativa al repo}."""
    import json

    p = _bibliografia_wiki_index_path(materia)
    if not p.is_file():
        return {}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}


def write_bibliografia_wiki_index(materia: str, index: dict[str, str]) -> None:
    import json

    _validate_materia(materia)
    _bibliografia_wiki_index_path(materia).write_text(
        json.dumps(index, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def read_bibliografia_global_wiki_index() -> dict[str, str]:
    import json

    p = _bibliografia_global_wiki_index_path()
    if not p.is_file():
        return {}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}


def write_bibliografia_global_wiki_index(index: dict[str, str]) -> None:
    import json

    _bibliografia_global_wiki_index_path().write_text(
        json.dumps(index, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def bibliography_source_rel_key(path: Path, *, materia: str | None) -> str:
    """Clave del índice: ruta relativa al árbol 01_raw/bibliografia/<materia>|global."""
    root = dir_bibliografia_global() if materia is None else dir_bibliografia_materia(materia)
    try:
        return path.relative_to(root).as_posix()
    except ValueError:
        return path.name


def pending_bibliografia_for_fichas(materia: str = DEFAULT_MATERIA) -> list[Path]:
    """Archivos en bibliografía de la materia sin ficha wiki generada."""
    idx = read_bibliografia_wiki_index(materia)
    out: list[Path] = []
    for p in list_bibliografia(materia):
        key = bibliography_source_rel_key(p, materia=materia)
        if key not in idx:
            out.append(p)
    return sorted(out, key=lambda x: x.as_posix().lower())


def pending_bibliografia_global_for_fichas() -> list[Path]:
    """Archivos en bibliografía global sin ficha wiki generada."""
    idx = read_bibliografia_global_wiki_index()
    out: list[Path] = []
    for p in list_bibliografia_global():
        key = bibliography_source_rel_key(p, materia=None)
        if key not in idx:
            out.append(p)
    return sorted(out, key=lambda x: x.as_posix().lower())


# ---------------------------------------------------------------------------
# Export helpers  (Markdown → .docx / .pdf)
# ---------------------------------------------------------------------------

import re

_HR_LINE = re.compile(r"^\s*([*\-_])\1{2,}\s*$")


def _export_md_is_hr_line(line: str) -> bool:
    """Línea solo con ---, ***, ___ (en Markdown no deben mostrarse en Word/PDF)."""
    return bool(_HR_LINE.match((line or "").strip()))


def _export_md_atx_heading(line: str) -> tuple[int, str] | None:
    """# … ###### título (ATX) con al menos un espacio tras #."""
    m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line.strip())
    if not m:
        return None
    level = min(len(m.group(1)), 6)
    title = m.group(2).strip()
    title = re.sub(r"\s+#+\s*$", "", title)  # cierre ATX: ### título ###
    return level, title


def _export_md_plain_segment(s: str) -> str:
    """Enlace [x](y), comillas, cursiva, negrita suelta: sin símbolos de edición a la vista."""
    t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"`([^`]+)`", r"\1", t)
    t = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"\1", t)
    return t


def _export_pdf_paragraph_markup(line: str) -> str:
    """Texto a marcado mínimo para reportlab: <b> y caracteres escapados."""
    from xml.sax.saxutils import escape

    t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line.strip())
    parts = re.split(r"(\*\*.+?\*\*)", t)
    buf: list[str] = []
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**") and len(p) > 4:
            inner = _export_md_plain_segment(p[2:-2])
            buf.append(f"<b>{escape(inner)}</b>")
        else:
            buf.append(escape(_export_md_plain_segment(p)))
    return "".join(buf)


def _export_docx_add_body_paragraph(doc, raw: str, _set_font) -> None:
    """Párrafo: **negrita** a runs; tramos sueltos sin * ` etc."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", raw.strip())
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    for part in re.split(r"(\*\*.+?\*\*)", t):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = p.add_run(_export_md_plain_segment(part[2:-2]))
            _set_font(r, 11, True)
        else:
            r = p.add_run(_export_md_plain_segment(part))
            _set_font(r, 11, False)


def export_to_docx(md_path: Path, metadata: dict | None = None) -> Path:
    """Convierte Markdown a Word en formato judicial real (Arial Narrow, márgenes de la Sala)."""
    from app.core.word_export import markdown_to_docx

    texto = md_path.read_text(encoding="utf-8", errors="replace")
    materia = infer_materia_from_resoluciones_md(md_path)
    out_dir = DIRS["exports"] / materia
    out_dir.mkdir(parents=True, exist_ok=True)
    dest = out_dir / md_path.with_suffix(".docx").name
    markdown_to_docx(texto, dest, metadata=metadata or {})
    return dest


def _export_to_docx_legacy(md_path: Path) -> Path:
    """Implementación anterior (Times New Roman genérico) — conservada por si se necesita."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    def _set_font(run, size, bold=False, color=None):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _add_heading(t: str, level: int) -> None:
        t2 = _export_md_plain_segment(t)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(t2)
        sizes = {1: 14, 2: 12, 3: 11, 4: 10, 5: 10, 6: 10}
        _set_font(r, sizes.get(min(level, 6), 11), True)
        p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
        p.paragraph_format.space_after = Pt(4)

    def _add_table(table_lines: list) -> None:
        rows = [
            l
            for l in table_lines
            if l.startswith("|") and not re.match(r"^\|[-| ]+\|$", l)
        ]
        if not rows:
            return
        cols = [c.strip() for c in rows[0].strip("|").split("|")]
        t = doc.add_table(rows=len(rows), cols=len(cols))
        t.style = "Table Grid"
        for r_idx, row_line in enumerate(rows):
            cells = [c.strip() for c in row_line.strip("|").split("|")]
            for c_idx, cell in enumerate(cells):
                if c_idx < len(t.rows[r_idx].cells):
                    c = t.rows[r_idx].cells[c_idx]
                    c.text = ""
                    r = c.paragraphs[0].add_run(_export_md_plain_segment(cell))
                    _set_font(r, 10, bold=(r_idx == 0))
        doc.add_paragraph()

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if _export_md_is_hr_line(line):
            i += 1
            continue
        h = _export_md_atx_heading(line)
        if h:
            _add_heading(h[1], h[0])
            i += 1
            continue
        if line.startswith("|"):
            tbl_lines: list = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            _add_table(tbl_lines)
            continue
        mnum = re.match(r"^(\d{1,3})\.\s+(.+)$", line)
        if mnum:
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(_export_md_plain_segment(mnum.group(2)))
            _set_font(r, 11, False)
            i += 1
            continue
        bq = re.match(r"^>\s?(.*)$", line)
        if bq:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(_export_md_plain_segment(bq.group(1)))
            _set_font(r, 11, False)
            i += 1
            continue
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(_export_md_plain_segment(line[2:]))
            _set_font(r, 11, False)
        elif not line.strip():
            pass
        else:
            _export_docx_add_body_paragraph(doc, line, _set_font)
        i += 1

    materia = infer_materia_from_resoluciones_md(md_path)
    out_dir = DIRS["exports"] / materia
    out_dir.mkdir(parents=True, exist_ok=True)
    dest = out_dir / md_path.with_suffix(".docx").name
    doc.save(dest)
    return dest


def export_to_pdf(md_path: Path) -> Path:
    """Markdown → PDF: sin reglas ---; títulos y listas sin mostrar # * en crudo."""
    from xml.sax.saxutils import escape

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors

    text = md_path.read_text(encoding="utf-8")
    materia = infer_materia_from_resoluciones_md(md_path)
    out_dir = DIRS["exports"] / materia
    out_dir.mkdir(parents=True, exist_ok=True)
    dest = out_dir / md_path.with_suffix(".pdf").name

    doc = SimpleDocTemplate(
        str(dest), pagesize=A4,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        leftMargin=3*cm, rightMargin=2.5*cm,
    )

    st_title  = ParagraphStyle("título",  fontName="Times-Bold",   fontSize=13, alignment=TA_CENTER, spaceAfter=6)
    st_h2     = ParagraphStyle("h2",      fontName="Times-Bold",   fontSize=11, alignment=TA_LEFT,   spaceBefore=8, spaceAfter=4)
    st_h3     = ParagraphStyle("h3",      fontName="Times-Bold",   fontSize=10, alignment=TA_LEFT,   spaceBefore=6, spaceAfter=3)
    st_h4     = ParagraphStyle("h4",      fontName="Times-Bold",   fontSize=9,  alignment=TA_LEFT,   spaceBefore=4, spaceAfter=2)
    st_h56    = ParagraphStyle("h56",     fontName="Times-Bold",   fontSize=9,  alignment=TA_LEFT,   spaceBefore=3, spaceAfter=2)
    st_body   = ParagraphStyle("cuerpo",  fontName="Times-Roman",  fontSize=10, alignment=TA_JUSTIFY, spaceAfter=4, leading=14)
    st_bullet = ParagraphStyle("bullet",  fontName="Times-Roman",  fontSize=10, alignment=TA_LEFT,   leftIndent=18, spaceAfter=2, leading=13)
    st_bq     = ParagraphStyle("cita",    fontName="Times-Roman",  fontSize=10, alignment=TA_JUSTIFY, leftIndent=18, spaceAfter=4, leading=14)
    # Tablas: celdas con Paragraph para que el texto ajuste al ancho (evita desborde)
    st_tbl_head = ParagraphStyle(
        "tblh",
        fontName="Times-Bold",
        fontSize=8,
        leading=10,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=0,
    )
    st_tbl_cell = ParagraphStyle(
        "tblc",
        fontName="Times-Roman",
        fontSize=8,
        leading=10,
        alignment=TA_JUSTIFY,
        spaceBefore=0,
        spaceAfter=0,
    )

    story = []
    lines = text.splitlines()
    i = 0

    def _heading_style(level: int) -> ParagraphStyle:
        if level == 1:
            return st_title
        if level == 2:
            return st_h2
        if level == 3:
            return st_h3
        if level == 4:
            return st_h4
        return st_h56

    while i < len(lines):
        line = lines[i]
        if _export_md_is_hr_line(line):
            i += 1
            continue
        h = _export_md_atx_heading(line)
        if h:
            lev, title = h
            story.append(
                Paragraph(escape(_export_md_plain_segment(title)), _heading_style(lev))
            )
            i += 1
            continue
        if line.startswith("|"):
            tbl_lines: list = []
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r"^\|[-| ]+\|$", lines[i]):
                    tbl_lines.append(lines[i])
                i += 1
            if tbl_lines:
                data = [
                    [_export_md_plain_segment(c.strip()) for c in r.strip("|").split("|")]
                    for r in tbl_lines
                ]
                col_count = max(len(r) for r in data)
                data = [r + [""] * (col_count - len(r)) for r in data]
                # Ancho útil = página menos márgenes (coincide con SimpleDocTemplate arriba)
                usable_w = A4[0] - 3 * cm - 2.5 * cm
                col_w = max(20, usable_w / col_count)
                data_flow: list = []
                for r_idx, row in enumerate(data):
                    row_p: list = []
                    is_head = r_idx == 0
                    for c in row:
                        cell_text = (c or "").replace("\n", " ").strip() or " "
                        if is_head:
                            ptxt = escape(_export_md_plain_segment(cell_text))
                            row_p.append(Paragraph(ptxt, st_tbl_head))
                        else:
                            row_p.append(Paragraph(_export_pdf_paragraph_markup(cell_text), st_tbl_cell))
                    data_flow.append(row_p)
                tbl = Table(
                    data_flow,
                    colWidths=[col_w] * col_count,
                    repeatRows=1,
                )
                tbl.setStyle(TableStyle([
                    ("GRID",        (0,0), (-1,-1), 0.4, colors.grey),
                    ("BACKGROUND",  (0,0), (-1,0),  colors.HexColor("#E8E8E8")),
                    ("VALIGN",      (0,0), (-1,-1), "TOP"),
                    ("LEFTPADDING",  (0,0), (-1,-1), 4),
                    ("RIGHTPADDING", (0,0), (-1,-1), 4),
                    ("TOPPADDING",  (0,0), (-1,-1), 3),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 3),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 6))
            continue
        mnum = re.match(r"^(\d{1,3})\.\s+(.+)$", line)
        if mnum:
            story.append(
                Paragraph(
                    f"{mnum.group(1)}. " + _export_pdf_paragraph_markup(mnum.group(2)),
                    st_bullet,
                )
            )
            i += 1
            continue
        bq = re.match(r"^>\s?(.*)$", line)
        if bq:
            story.append(Paragraph(_export_pdf_paragraph_markup(bq.group(1)), st_bq))
            i += 1
            continue
        if line.startswith("- ") or line.startswith("* "):
            story.append(Paragraph("• " + _export_pdf_paragraph_markup(line[2:].strip()), st_bullet))
        elif not line.strip():
            story.append(Spacer(1, 3))
        else:
            story.append(Paragraph(_export_pdf_paragraph_markup(line), st_body))
        i += 1

    doc.build(story)
    return dest




# ---------------------------------------------------------------------------
# Prompt builder
# ---------------------------------------------------------------------------

def _list_estilo_files(materia: str = DEFAULT_MATERIA) -> list[Path]:
    """.md en plantillas/<materia>/estilo/ o, por compatibilidad, plantillas/estilo/."""
    candidates = [
        DIRS["plantillas"] / materia / "estilo",
        DIRS["plantillas"] / "estilo",
    ]
    for estilo_dir in candidates:
        if estilo_dir.is_dir():
            return sorted(
                [
                    f
                    for f in estilo_dir.iterdir()
                    if f.is_file()
                    and f.suffix.lower() == ".md"
                    and not f.name.startswith(".")
                ],
                key=lambda p: p.name.lower(),
            )
    return []


def _resolve_plantilla_path(plantilla_name: str, materia: str = DEFAULT_MATERIA) -> str:
    """Resuelve la ruta de una plantilla (prioriza la carpeta de la materia)."""
    _validate_materia(materia)
    for root in (DIRS["plantillas"] / materia, DIRS["plantillas"]):
        matches = list(root.rglob(plantilla_name))
        if matches:
            return str(matches[0].relative_to(BASE_DIR))
    return f"01_raw/plantillas/{materia}/{plantilla_name}"


_SEP = "═" * 55


def _block_rol() -> str:
    return (
        f"{_SEP}\n"
        "BLOQUE 1 · ROL Y RESTRICCIONES\n"
        f"{_SEP}\n"
        "Eres asistente jurídico del juez de la Sala Penal de Apelaciones.\n"
        "• No inventes casaciones, artículos, libros ni referencias.\n"
        "• Si no hay respaldo documentado en el repositorio, dilo explícitamente.\n"
        "• Usa ÚNICAMENTE las fuentes y materiales proporcionados en los bloques siguientes.\n"
        "• Estilo: técnico-jurídico humanizado, preciso, elocuente sin retórica vacía.\n"
        + BLOQUE_ROL_ANTIAI_UNA_LINEA_ES
        + "• Respeta el razonamiento silogístico judicial.\n"
    )


def _block_parametros_institucionales_peru() -> str:
    """Parámetros estables de redacción judicial fijados por magistratura."""
    return (
        f"{_SEP}\n"
        "BLOQUE 1B · PARÁMETROS INSTITUCIONALES DE REDACCIÓN (PERÚ)\n"
        f"{_SEP}\n"
        "Estas reglas generales se aplican en TODOS los casos, además de la plantilla, "
        "la instrucción general de materia y la instrucción particular del expediente.\n\n"
        "Contexto fijo del acto:\n"
        "• Tipo: resolución judicial.\n"
        "• Ámbito: penal y constitucional.\n"
        "• País/ordenamiento: Perú.\n"
        "• Instancia: apelación.\n\n"
        "Estructura y forma:\n"
        "• Mantén estructura Vistos, Considerando(s) y Resuelve.\n"
        "• Considerandos numerados.\n"
        "• Sin tope rígido de extensión: prioriza brevedad, claridad, precisión y profundidad,\n"
        "  sin relleno ni párrafos protocolares vacíos.\n\n"
        "Estilo y tratamiento:\n"
        "• Lenguaje claro sin perder rigor técnico-jurídico.\n"
        "• Redacción en tercera persona impersonal.\n"
        "• Identifica a las partes por nombre completo, salvo en delitos contra libertad sexual:\n"
        "  víctima (y menores cuando corresponda) con iniciales.\n\n"
        "Redacción judicial clara (operativa):\n"
        "• Usa frases preferentemente breves y párrafos de una sola idea.\n"
        "• Evita fórmulas vacías o de relleno (por ejemplo: cabe señalar, es menester,\n"
        "  en ese orden de ideas) salvo necesidad argumentativa real.\n"
        "• No repitas doctrina, jurisprudencia ni hechos si no aportan al punto controvertido.\n"
        "• Cada considerando debe cumplir una función verificable: hecho probado, norma aplicable,\n"
        "  estándar jurisprudencial, inferencia del caso o respuesta a un agravio.\n"
        "• Conserva la terminología penal, procesal penal y constitucional necesaria, pero explica\n"
        "  su aplicación al caso con lenguaje directo.\n\n"
        "Modo de síntesis jurídica (regla por defecto):\n"
        "• Redacta de modo breve, claro y preciso, sin sacrificar suficiencia jurídica.\n"
        "• No sacrifiques fundamentación por brevedad; sacrifica redundancia, retórica y\n"
        "  transcripción innecesaria, **excepto en cargos imputados por el MP, que van en transcripción literal**.\n"
        "• En cada punto relevante articula: hecho o agravio, norma/estándar, valoración e inferencia,\n"
        "  y conclusión.\n\n"
        "Jerarquía de fundamentación (triple anclaje):\n"
        "• Orden de peso: normativo > jurisprudencial > doctrinal.\n"
        "• Eje argumental: Constitución, tratados, ley y reglamento; jurisprudencia como criterio\n"
        "  de interpretación/aplicación; doctrina en rol auxiliar.\n"
        "• En hábeas corpus prioriza: (i) TC Perú, (ii) Corte Suprema penal (AP/Casaciones),\n"
        "  (iii) Corte IDH para control de convencionalidad, (iv) doctrina nacional.\n"
        "• Cuando corresponda, considera normativa convencional y decisiones de la Corte IDH/CIDH.\n"
        "• Cuando corresponda, aplica control de motivación con STC 00728-2008-PHC/TC\n"
        "  (caso Giuliana Llamoja) y STC Exp. N.° 01747-2013-PA/TC.\n"
        "• Si el caso lo exige, aplica test de proporcionalidad (idoneidad, necesidad,\n"
        "  proporcionalidad en sentido estricto) con enfoque operativo alineado al TC.\n\n"
        "Jurisprudencia y verificación:\n"
        "• Nunca inventes fuentes. Si no hay respaldo verificable, decláralo expresamente.\n"
        "• Formatos de cita válidos (referenciales):\n"
        "  - STC Exp. N.° XXXX-AAAA-XX/TC, fundamento N.\n"
        "  - Casación N.° XXX-AAAA, sede, considerando N.\n"
        "  - Doctrina en APA + nota al pie cuando corresponda.\n\n"
        "Operativa de redacción:\n"
        "• Si el expediente es extenso, inicia con mapeo interno de puntos controvertidos\n"
        "  antes de redactar el texto final.\n"
        "• Si detectas que el caso amerita test de proporcionalidad o control de motivación tipo\n"
        "  Llamoja/STC 01747-2013-PA/TC, incorpora ese enfoque de modo expreso en el desarrollo.\n"
        "• Los datos del expediente no son parte del razonamiento ni del cuerpo resolutivo: deben ir\n"
        "  en el encabezado previo al rótulo «Auto de Vista» o «Sentencia de Vista», según corresponda.\n"
        "• Ese encabezado previo debe incluir, en todos los casos con la misma estructura de rubros y cuando\n"
        "  conste en las fuentes: expediente, imputado(s), delito, agraviado y procedencia.\n"
        "  Si falta un rubro, extráelos de la documentación del expediente o del formulario; solo marca\n"
        "  como no determinado un rubro que no figure en las fuentes.\n"
        "• Si el magistrado lo pide, además del borrador en el sentido indicado, propone\n"
        "  alternativas razonadas de fallo (confirmar/revocar/nulidad) cuando el expediente lo permita.\n"
        "• La postura indicada por el magistrado es vinculante para el proyecto principal: si se ordena\n"
        "  CONFIRMAR, el razonamiento y la parte resolutiva deben confirmar la resolución apelada;\n"
        "  queda prohibido revocarla, reformarla o sustituir la medida por comparecencia, salvo orden\n"
        "  expresa posterior del magistrado.\n"
        "• El rubro de fundamentos del recurso de apelación debe ser estrictamente descriptivo: resume\n"
        "  solo agravios, expresiones y argumentos que consten en el escrito de apelación. No agregues\n"
        "  argumentos plausibles, defensas genéricas ni inferencias no formuladas por el recurrente.\n"
        "• La absolución de agravios se hace únicamente sobre agravios surgidos del recurso de apelación.\n"
        "  No conviertas en agravios los alegatos finales de primera instancia, teorías defensivas del juicio\n"
        "  ni argumentos consignados en la sentencia apelada, salvo que el recurso de apelación los reproduzca\n"
        "  o los incorpore expresamente como agravio.\n"
        "• Antes de entregar, realiza una depuración interna: elimina redundancias, fusiona párrafos\n"
        "  que repiten la misma idea, suprime introducciones protocolares sin contenido, verifica que\n"
        "  cada considerando responda a un problema jurídico o agravio, y conserva solo citas normativas\n"
        "  y jurisprudenciales necesarias para resolver.\n\n"
        f"{CARGOS_MP_LITERAL_ES}\n"
    )


def _block_instruccion_permanente(instruccion_general: str, materia: str) -> str:
    ig = (instruccion_general or "").strip()
    if not ig:
        return ""
    return (
        f"{_SEP}\n"
        f"BLOQUE 2 · INSTRUCCIÓN PERMANENTE DE LA MATERIA ({materia_label(materia)})\n"
        f"{_SEP}\n"
        f"{ig}\n"
    )


def _block_plantilla(plantilla_name: str, materia: str) -> str:
    if not plantilla_name or plantilla_name == "sin_plantilla":
        return ""
    plantilla_path = _resolve_plantilla_path(plantilla_name, materia)
    return (
        f"{_SEP}\n"
        "BLOQUE 3 · PLANTILLA (MOLDE ESTRUCTURAL)\n"
        f"{_SEP}\n"
        f"Lee PRIMERO: `{plantilla_path}`\n"
        "Sigue EXACTAMENTE su estructura: encabezado, numeración, secciones,\n"
        "tabla de datos, orden de considerandos y formato de firma S.S.\n"
        "La parte común de la plantilla (lugar y fecha, Vistos, rótulos, orden de secciones,\n"
        "formato de encabezado y cierre) debe seguirse en cuanto no se oponga a la particularidad\n"
        "del caso. La plantilla se añade para seguir formato e ítems a desarrollar; NO inventes\n"
        "estructura propia ni omitas secciones de la plantilla.\n"
    )


def _block_fuentes(
    folder_name: str,
    slots: dict[str, list[str]] | None,
    materia: str,
) -> str:
    header = (
        f"{_SEP}\n"
        "BLOQUE 4 · FUENTES DEL EXPEDIENTE\n"
        f"{_SEP}\n"
        "Estos son los hechos del caso concreto. Analízalos, numera los actos\n"
        "de investigación y cita textualmente los pasajes de respaldo cuando\n"
        "fundamentes conclusiones.\n"
        + rules_for_untrusted_sources()
    )

    # Sin slots tipificados: fallback al comportamiento clásico
    if not slots or not any(slots.values()):
        return header + f"\nLee todos los archivos en `01_raw/{folder_name}/`\n"

    labels = slot_labels_for(materia)
    parts: list[str] = []
    for key in SLOT_KEYS:
        paths = slots.get(key) or []
        if not paths:
            continue
        label = labels.get(key, key.replace("_", " ").capitalize())
        parts.append(f"\n### {label}")
        for p in paths:
            parts.append(f"  - `{p}`")
    return header + "\n".join(parts) + "\n"


def _block_bibliografia(materia: str, bibliografia_activa: list[str] | None) -> str:
    bib_rel = f"01_raw/bibliografia/{materia}"
    if not bibliografia_activa:
        return (
            f"{_SEP}\n"
            "BLOQUE 5 · BIBLIOGRAFÍA DISPONIBLE (consulta opcional)\n"
            f"{_SEP}\n"
            f"Disponible en `{bib_rel}/`. Consulta solo si es directamente pertinente\n"
            "al caso concreto. NO vuelques doctrina no solicitada ni llenes la\n"
            "resolución con citas dogmáticas de adorno.\n"
            "Test de pertinencia obligatorio: la ratio del precedente debe coincidir\n"
            "con el punto jurídico que fundamentas; si solo trata genéricamente la\n"
            "materia o un tema distinto, omite la cita.\n"
            + rules_for_untrusted_sources()
        )
    items = "\n".join(f"  - `{p}`" for p in bibliografia_activa)
    return (
        f"{_SEP}\n"
        "BLOQUE 5 · BIBLIOGRAFÍA ACTIVA (seleccionada por el magistrado)\n"
        f"{_SEP}\n"
        f"Material de consulta marcado como relevante para este caso:\n{items}\n"
        "Directiva: cita solo si es directamente pertinente. Calidad sobre cantidad.\n"
        "Test de pertinencia obligatorio antes de citar: (1) mismo problema jurídico,\n"
        "(2) misma ratio aplicable al punto que fundamentas, y (3) analogía fáctica no engañosa.\n"
        "Si no supera ese test, omite la cita aunque el archivo esté listado como bibliografía activa.\n"
        "No uses una casación real para una regla distinta a su ratio: por ejemplo, una casación sobre\n"
        "interés superior del niño no sustenta salud, adicciones, atención penitenciaria ni otro tema ajeno.\n"
        + rules_for_untrusted_sources()
    )


def _block_estilo(materia: str, corpus_sample: list[str] | None) -> str:
    lines: list[str] = []
    any_content = False

    estilo_files = _list_estilo_files(materia)
    if estilo_files:
        any_content = True
        try:
            rel0 = estilo_files[0].relative_to(BASE_DIR)
            estilo_base = str(rel0.parent).replace("\\", "/")
        except ValueError:
            estilo_base = f"01_raw/plantillas/{materia}/estilo"
        lines.append("\nTextos de referencia estilística:")
        for f in estilo_files:
            lines.append(f"  - `{estilo_base}/{f.name}`")

    if corpus_sample:
        any_content = True
        lines.append(
            "\nResoluciones previas firmadas del magistrado en casos similares\n"
            "(fuente principal de estilo — imita voz, sintaxis y registro):"
        )
        for p in corpus_sample:
            lines.append(f"  - `{p}`")

    if not any_content:
        return ""

    return (
        f"{_SEP}\n"
        "BLOQUE 6 · REFERENCIA DE ESTILO DEL MAGISTRADO\n"
        f"{_SEP}\n"
        + "\n".join(lines)
        + "\n\nAdopta la voz, construcciones sintácticas y registro argumentativo.\n"
        "NO uses fórmulas genéricas de IA ni retórica hueca.\n"
    )


def _formato_postura(postura: str, postura_personalizada: str) -> str:
    p = (postura or "").strip().lower()
    if p == "confirmar" or "confir" in p:
        return (
            "CONFIRMAR la resolución apelada. Valida el razonamiento del a quo;\n"
            "desestima cada agravio argumentativamente; refuerza con bibliografía\n"
            "pertinente SOLO si aplica al caso. La parte resolutiva debe confirmar;\n"
            "no revoques, no reformes ni sustituyas la medida por comparecencia u otra consecuencia incompatible."
        )
    if p == "revocar" or "revo" in p:
        return (
            "REVOCAR la resolución apelada. Identifica los errores del a quo y\n"
            "construye el razonamiento inverso. Cada presupuesto que el a quo\n"
            "consideró cumplido debe rebatirse con argumentos concretos de las fuentes."
        )
    if p == "nulificar" or "nulif" in p or "nulidad" in p:
        return (
            "NULIFICAR la resolución apelada. Identifica el vicio procesal concreto\n"
            "(nulidad absoluta o relativa), cita el dispositivo legal que sustenta\n"
            "la nulidad y delimita los efectos (qué se reenvía y a qué estadio)."
        )
    extra = (postura_personalizada or "").strip()
    if p == "otros" or extra:
        base = "OTROS — postura e instrucciones definidas por el magistrado:"
        return f"{base}\n{extra}" if extra else base
    return f"{postura}. Determina la postura al leer las fuentes y justifícala."


def _block_configuracion(
    expediente: str,
    imputados: str,
    delito: str,
    agraviado: str,
    juzgado: str,
    postura: str,
    postura_personalizada: str,
    agravios: str,
) -> str:
    meta_lines = []
    if expediente: meta_lines.append(f"- Expediente: {expediente}")
    if imputados:  meta_lines.append(f"- Imputado(s): {imputados}")
    if delito:     meta_lines.append(f"- Delito: {delito}")
    if agraviado:  meta_lines.append(f"- Agraviado: {agraviado}")
    if juzgado:    meta_lines.append(f"- Juzgado de origen: {juzgado}")
    if not meta_lines:
        meta_lines.append("(Extrae los metadatos directamente de las fuentes del caso.)")

    ag = (agravios or "").strip()
    if ag:
        agravios_block = f"\nAgravios a absolver (uno por párrafo):\n{ag}\n"
    else:
        agravios_block = (
            "\nAgravios: extráelos del recurso de apelación, numéralos y absuelve\n"
            "cada uno en párrafo propio — ninguno puede omitirse ni fusionarse. En el rubro\n"
            "fundamentos del recurso, no atribuyas al recurrente expresiones o argumentos\n"
            "que no consten en el escrito de apelación. No uses alegatos finales, defensas\n"
            "de primera instancia ni argumentos recogidos en la sentencia como agravios si\n"
            "el recurso no los reproduce.\n"
        )

    return (
        f"{_SEP}\n"
        "BLOQUE 7 · CONFIGURACIÓN DEL CASO\n"
        f"{_SEP}\n"
        "Metadatos:\n"
        + "\n".join(meta_lines)
        + "\n\nPostura:\n"
        + _formato_postura(postura, postura_personalizada)
        + "\n"
        + agravios_block
    )


def _block_tarea_generar(tipo: str, materia: str, caso_num: str) -> str:
    return (
        f"{_SEP}\n"
        "BLOQUE 8 · TAREA — GENERAR DESDE CERO\n"
        f"{_SEP}\n"
        f"Redacta la **{tipo}** completa, lista para revisión y firma:\n"
        "1. Sigue EXACTAMENTE la estructura de la plantilla (Bloque 3), en especial lugar y fecha,\n"
        "   Vistos, rótulos, orden de secciones, formato de encabezado, cierre y firma, salvo aquello\n"
        "   que contradiga los datos o particularidades del caso.\n"
        "2. Adopta el estilo del magistrado (Bloque 6).\n"
        "3. Aplica la postura indicada (Bloque 7) — sin oscilar.\n"
        "4. Si la postura es CONFIRMAR, no redactes fallo revocatorio ni sustituyas la medida por\n"
        "   comparecencia u otra consecuencia incompatible.\n"
        "5. Antes del rótulo «Auto de Vista» o «Sentencia de Vista», coloca el encabezado con los\n"
        "   datos (expediente, imputado(s), delito, agraviado y procedencia),\n"
        "   extrayéndolos de las fuentes si faltan; no los insertes como parte del cuerpo resolutivo.\n"
        "6. En fundamentos del recurso, resume solo lo que aparece en el recurso de apelación; no extraigas\n"
        "   agravios de alegatos finales de primera instancia ni de la sentencia apelada si el recurso no los replantea.\n"
        "7. Absuelve cada agravio real en párrafo propio.\n"
        "8. Si el expediente es voluminoso, mapea primero puntos controvertidos y luego redacta.\n"
        "9. Si hay base razonable en actuados y se te pide, sugiere alternativas de fallo\n"
        "   (confirmar/revocar/nulidad) de forma comparada y fundamentada.\n"
        "10. Guarda el borrador en "
        f"`03_outputs/resoluciones/{materia}/` como versión nueva\n"
        f"   (incluye `{caso_num}` en el nombre; NO sobreescribas versiones previas).\n"
    )


def _block_tarea_continuar(borrador_path: str) -> str:
    return (
        f"{_SEP}\n"
        "BLOQUE 8 · TAREA — CONTINUAR BORRADOR DEL MAGISTRADO\n"
        f"{_SEP}\n"
        f"El magistrado inició la resolución en `{borrador_path}`.\n"
        "El archivo puede ser **Markdown (.md), Word (.doc/.docx), PDF (.pdf) o Pages (.pages)**.\n"
        "Ábrelo y extrae o lee el texto con las herramientas disponibles en el entorno.\n"
        "• **PDF / Word:** no inventes texto que no conste en el documento.\n"
        "• **Pages:** si no puedes leer el paquete .pages, pide exportar a Word o PDF y usa esa ruta.\n\n"
        "Busca el marcador `[[CONTINUAR AQUÍ]]`. Si no existe, continúa\n"
        "después del último párrafo escrito.\n\n"
        "Directivas estrictas:\n"
        "1. NO reescribas, parafrasees ni resumas lo que el magistrado ya escribió.\n"
        "2. Continúa EXACTAMENTE el hilo argumentativo ya desarrollado.\n"
        "3. El texto existente es la MEJOR referencia de estilo disponible —\n"
        "   imita su voz, tono y nivel de elocuencia (prima sobre el Bloque 6).\n"
        "4. Si el Bloque 7 no fija postura, dedúcela de lo ya escrito.\n"
        "5. Completa las secciones faltantes de la plantilla (Bloque 3) manteniendo\n"
        "   coherencia con hechos y doctrina ya citados por el magistrado.\n"
        "6. Al terminar, guarda el resultado completo como **versión nueva** (idealmente `.md` o `.docx`\n"
        "   en `03_outputs/resoluciones/<materia>/`) sin sobreescribir el archivo original.\n"
    )


def _block_reglas_transversales(materia: str) -> str:
    bib_rel = f"01_raw/bibliografia/{materia}"
    return (
        f"{_SEP}\n"
        "REGLAS TRANSVERSALES\n"
        f"{_SEP}\n"
        "Citas jurisprudenciales (CRÍTICO):\n"
        f"• Usa ÚNICAMENTE jurisprudencia que exista en `{bib_rel}/`, en\n"
        "  `02_wiki/jurisprudencia/` o en los archivos del caso.\n"
        "• Si necesitas una casación/STC que no esté documentada, escribe\n"
        "  `[CITA PENDIENTE DE VERIFICACIÓN]` en vez de inventar el número.\n"
        "• NUNCA inventes números de casación, folios, fechas ni nombres.\n\n"
        "Estructura funcional a replicar (orden — adaptar rótulos al Bloque 3):\n"
        "1) Encabezado y fórmula de apertura: datos del expediente, partes, procedencia,\n"
        "   identificación de la resolución y fórmula de «Vistos».\n"
        "2) Planteamiento del caso: resolución impugnada; recurso de la defensa (agravios);\n"
        "   intervención oral de la defensa; posición del Ministerio Público; intervención del actor civil;\n"
        "   interrogatorio y autodefensa del acusado; problema jurídico.\n"
        "3) Cargos imputados: **transcripción literal** del escrito fiscal (ranura solicitud_inicial): "
        "hechos imputados, circunstancias y calificación penal.\n"
        "4) Razonamiento (considerandos): núcleo argumentativo.\n"
        "5) Costas — cuando corresponda.\n"
        "6) Decisión (Resuelve).\n\n"
        "Profundidad argumentativa por función (no por número romano de plantilla):\n"
        "• Encabezado/apertura: registro descriptivo; **Cargos imputados del MP: transcripción literal** "
        "del escrito fiscal en solicitud_inicial, sin valoración ni anticipo de conclusiones.\n"
        "• Planteamiento del caso:\n"
        "  · Resolución impugnada: la parte dispositiva/resolutiva de la resolución apelada.\n"
        "  · Agravios de la defensa: reproducción mediante **resumen fidedigno y fiel** al recurso de apelación\n"
        "    — sin parafrasear de modo que se altere el sentido ni fusionar agravios distintos: cada agravio\n"
        "    conserva identidad y orden. No incorporar alegatos finales ni defensas de primera instancia que\n"
        "    no hayan sido reproducidos como agravios en el recurso. No añadir argumentos no planteados por la parte.\n"
        "  · Demás voces procesales (Ministerio Público, actor civil o agraviado, interrogatorio y autodefensa):\n"
        "    síntesis fiel, rigor técnico-legal, respetando sentido y terminología originales; sin reproducción\n"
        "    literal ni añadidos.\n"
        "• Razonamiento (considerandos): **máxima elaboración** — centro de gravedad. Desarrollar marco de\n"
        "  competencia revisora, premisas normativas, tipificación aplicable con **descripción literal del tipo**,\n"
        "  estándar de prueba y análisis del caso concreto (el **análisis individual de los medios probatorios**\n"
        "  en forma de listado efectuado por el juez de primera instancia, como la **valoración en conjunto**),\n"
        "  y absolver los agravios. Preferencia: **un considerando por cada agravio**, mismo orden de planteamiento;\n"
        "  si agravios conexos se tratan juntos, indícalo explícitamente por claridad.\n"
        "• Decisión (Resuelve): **sintética**, sin razonamiento nuevo; solo lo resuelto, en concordancia estricta\n"
        "  con los considerandos.\n"
        "Si la plantilla agrupa funciones en un solo apartado, obedece la plantilla pero conserva esta gradación interna.\n"
        "• Redacción en tercera persona impersonal y considerandos numerados.\n"
    )


def _block_reglas_pp_calificacion() -> str:
    return (
        "Complemento PP — calificación receptación:\n"
        "• Tras la transcripción literal fiscal, la calificación debe incluir tipicidad completa "
        "(art. 194 CP + agravante art. 195 + art. 427 CP u otros), no solo el párrafo de pena.\n"
    )


def _block_checklist_pp() -> str:
    """Checklist específico para resoluciones de prisión preventiva."""
    return (
        "Checklist (no omitir ningún punto):\n"
        "- [ ] Fecha completa (día, mes, año) en el encabezado\n"
        "- [ ] Tabla de datos: expediente, imputados, delito, agraviado, procedencia\n"
        "- [ ] Sección I: resolución objeto del recurso\n"
        "- [ ] Sección II: cargos imputados — transcripción literal del escrito fiscal (solicitud_inicial)\n"
        "- [ ] Sección II: calificación jurídica completa (tipo base + agravante + delitos concurrentes)\n"
        "- [ ] Sección III: agravios de la defensa (TODOS)\n"
        "- [ ] Sección IV: marco jurídico\n"
        "- [ ] Sección V: análisis de presupuestos\n"
        "    · Fumus boni iuris — cada elemento de convicción por separado\n"
        "    · Prognosis de pena — fundamento en el AP correspondiente\n"
        "    · Peligro de fuga (art. 269 CPP) — criterio por criterio\n"
        "    · Peligro de obstaculización (art. 270 CPP) — criterio por criterio\n"
        "    · Test de proporcionalidad — idoneidad, necesidad, proporcionalidad estricta\n"
        "    · Plazo — justificación cronológica con diligencias pendientes concretas\n"
        "- [ ] Sección VI: decisión con numerales PRIMERO, SEGUNDO…\n"
        "- [ ] Firma S.S. con nombres de los tres magistrados\n"
    )


def _block_checklist_generico() -> str:
    return (
        "Checklist (no omitir ningún punto):\n"
        "- [ ] Orden funcional según REGLAS TRANSVERSALES (apertura → planteamiento → cargos → considerandos\n"
        "      → costas si hay → Resuelve); rótulos del Bloque 3 prevalecen pero respeta esa secuencia lógica.\n"
        "- [ ] Fecha completa (día, mes, año) en el encabezado\n"
        "- [ ] Tabla de datos del expediente\n"
        "- [ ] Antecedentes / hechos relevantes\n"
        "- [ ] Reproducción íntegra de los agravios\n"
        "- [ ] Marco jurídico aplicable\n"
        "- [ ] Análisis fundamentado punto por punto\n"
        "- [ ] Decisión clara y numerada\n"
        "- [ ] Firma S.S.\n"
    )


def _block_wiki_workflow(caso_num: str) -> str:
    return (
        f"{_SEP}\n"
        "FLUJO WIKI · EJECUTAR OBLIGATORIAMENTE DESPUÉS DE REDACTAR\n"
        f"{_SEP}\n"
        "A) Nota del caso\n"
        f"   Crea/actualiza `02_wiki/casos/Caso{caso_num}.md`: resumen 3 puntos,\n"
        "   datos, hechos, agravios, postura, conceptos aplicados, jurisprudencia\n"
        "   usada y casos similares.\n\n"
        "B) Notas atómicas de conceptos\n"
        "   Por cada concepto jurídico relevante aplicado, crea/actualiza\n"
        "   `02_wiki/conceptos/NombreConcepto.md` y conecta con [[WikiLink]].\n"
        "   Si ya existe, agrega este caso en 'Casos donde se aplicó'.\n\n"
        "C) Notas de jurisprudencia\n"
        "   Por cada casación/AP/STC citada, crea/actualiza `02_wiki/jurisprudencia/`.\n"
        "   Si ya existe, agrega este caso en 'Casos donde se usó'.\n\n"
        "D) Índice maestro\n"
        "   Actualiza `02_wiki/INDEX.md` con el caso + conceptos y jurisprudencia nuevos.\n\n"
        "E) Reporte final\n"
        "   Lista todo lo creado/actualizado: resolución, nota de caso, conceptos,\n"
        "   jurisprudencia y conexiones nuevas identificadas.\n"
    )


def build_cursor_prompt(
    folder_name: str,
    plantilla_name: str,
    tipo: str = "Auto de Vista — Apelación PP",
    postura: str = "Por determinar",
    instruccion: str = "",
    instruccion_general: str = "",
    expediente: str = "",
    imputados: str = "",
    delito: str = "",
    agraviado: str = "",
    juzgado: str = "",
    materia: str = DEFAULT_MATERIA,
    # --- NUEVOS PARÁMETROS (opcionales, compatibles con llamadas anteriores) ---
    postura_personalizada: str = "",
    agravios: str = "",
    modo: str = "generar",                                  # "generar" | "continuar"
    borrador_path: str = "",
    slots: dict[str, list[str]] | None = None,              # {slot_key: [rutas rel.]}
    bibliografia_activa: list[str] | None = None,           # rutas rel. seleccionadas
    corpus_sample: list[str] | None = None,                 # resoluciones previas sample
) -> str:
    """Arma el prompt maestro para Cursor/Claude Code con 8 bloques estructurados.

    Backward-compatible: si no se pasan los nuevos parámetros (`slots`,
    `bibliografia_activa`, `corpus_sample`, `modo`), funciona igual que la
    versión anterior pero con la nueva arquitectura de bloques.
    """
    _validate_materia(materia)
    caso_num = caso_num_from_folder_key(folder_name)

    # Tarea (depende del modo)
    if modo == "continuar" and borrador_path:
        tarea_block = _block_tarea_continuar(borrador_path)
    else:
        tarea_block = _block_tarea_generar(tipo, materia, caso_num)

    # Checklist específico para PP; genérico para el resto
    checklist = (
        _block_checklist_pp()
        if materia == MATERIA_PRISION_PREVENTIVA
        else _block_checklist_generico()
    )

    # Instrucción particular del caso (iterable, opcional)
    ip = (instruccion or "").strip()
    ip_block = ""
    if ip:
        ip_block = (
            f"{_SEP}\n"
            "INSTRUCCIÓN PARTICULAR DE ESTE CASO (iterable)\n"
            f"{_SEP}\n"
            f"{ip}\n"
        )

    blocks = [
        f"# Procesar caso: {folder_name}\n",
        _block_rol(),
        _block_parametros_institucionales_peru(),
        _block_instruccion_permanente(instruccion_general, materia),
        _block_plantilla(plantilla_name, materia),
        _block_fuentes(folder_name, slots, materia),
        _block_bibliografia(materia, bibliografia_activa),
        _block_estilo(materia, corpus_sample),
        _block_configuracion(
            expediente, imputados, delito, agraviado, juzgado,
            postura, postura_personalizada, agravios,
        ),
        tarea_block,
        _block_reglas_transversales(materia)
        + (
            "\n" + _block_reglas_pp_calificacion()
            if materia == MATERIA_PRISION_PREVENTIVA
            else ""
        )
        + "\n"
        + checklist,
        _block_wiki_workflow(caso_num),
        ip_block,
    ]

    return "\n".join(b for b in blocks if b)
